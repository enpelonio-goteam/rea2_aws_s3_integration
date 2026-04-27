from fastapi import FastAPI, HTTPException, UploadFile, File, Header
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import boto3
import uuid
from datetime import datetime
import os
from typing import Optional
import json
import logging
from dotenv import load_dotenv
import requests
import base64
import mimetypes
import io
import tempfile
import time
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse, parse_qs, unquote
from mutagen.mp3 import MP3
from io import BytesIO
import loomdl
import assemblyai as aai
from PIL import Image, ImageOps

# Load environment variables from .env file
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Pydantic model for request body
class HTMLUploadRequest(BaseModel):
    html_content: str
    filename: Optional[str] = None
    content_type: str = "text/html"

class ImageGenerationRequest(BaseModel):
    prompt: str
    size: Optional[str] = "1024x1024"
    quality: Optional[str] = "auto"

class LoomVideoRequest(BaseModel):
    loom_url: str
    filename: Optional[str] = None

class LinkedInVideoRequest(BaseModel):
    video_url: str  # URL of the video to upload
    caption: str
    visibility: Optional[str] = "PUBLIC"  # PUBLIC, CONNECTIONS, or LOGGED_IN
    filename: Optional[str] = None

class HTMLToImageRequest(BaseModel):
    html_content: str
    width: Optional[int] = 1080  # Default Instagram post width
    height: Optional[int] = 1080  # Default Instagram post height (square)
    filename: Optional[str] = None
    quality: Optional[int] = 95  # PNG quality (0-100)

class UrlToGoogleDriveRequest(BaseModel):
    file_url: str  # Publicly accessible URL to download from
    folder_id: str  # Google Drive folder ID to upload to
    filename: Optional[str] = None  # Optional filename for the uploaded file

class ElevenLabsSpeechRequest(BaseModel):
    """
    Request model for Eleven Labs text-to-speech generation.
    Based on: https://elevenlabs.io/docs/api-reference/text-to-speech/convert
    """
    voice_id: str  # Eleven Labs voice ID to use for speech generation
    text: str  # Text to convert to speech
    model_id: Optional[str] = "eleven_multilingual_v2"  # Model ID (default: eleven_multilingual_v2)
    filename: Optional[str] = None  # Optional filename for the uploaded audio
    # Voice settings (all optional - uses Eleven Labs defaults if not provided)
    stability: Optional[float] = None  # 0.0-1.0, default 0.5. Lower = more emotional range, higher = more stable/monotonous
    similarity_boost: Optional[float] = None  # 0.0-1.0, default 0.75. How closely AI adheres to original voice
    style: Optional[float] = None  # 0.0-1.0, default 0. Style exaggeration (higher = more computational resources)
    use_speaker_boost: Optional[bool] = None  # default true. Boosts similarity to original speaker (increases latency)
    speed: Optional[float] = None  # default 1.0. Speech speed (<1.0 slower, >1.0 faster)

class HeyGenAvatarIVRequest(BaseModel):
    """
    Request model for HeyGen Avatar IV video generation.
    Based on: https://docs.heygen.com/reference/create-an-avatar-video-v2
    
    Note: Instead of image_key, this endpoint accepts image_url which will be 
    automatically uploaded to HeyGen using their Upload Asset API.
    
    You can either:
    1. Use TTS (text-to-speech): Provide script and voice_id
    2. Use pre-recorded audio: Provide audio_url or audio_asset_id (script and voice_id not required)
    """
    image_url: str  # Publicly accessible URL of the image to use for avatar (will be uploaded to HeyGen)
    script: Optional[str] = None  # The text that the avatar will speak (required for TTS, optional if using audio_url)
    voice_id: Optional[str] = None  # The voice ID to use for speech generation (required for TTS, optional if using audio_url)
    video_title: Optional[str] = None  # Optional title for the video
    video_orientation: Optional[str] = None  # Video orientation (e.g., "landscape", "portrait", "square")
    fit: Optional[str] = None  # How the image fits in the video frame: "cover" or "contain"
    custom_motion_prompt: Optional[str] = None  # Custom motion prompt for avatar movements
    enhance_custom_motion_prompt: Optional[bool] = None  # Whether to enhance the custom motion prompt with AI
    audio_url: Optional[str] = None  # URL of an audio file to use instead of TTS
    audio_asset_id: Optional[str] = None  # Asset ID of a previously uploaded audio file
    super_resolution: Optional[bool] = True  # Enhance talking photo image quality

class GoogleDriveToS3Request(BaseModel):
    """
    Request model for downloading a file from Google Drive and uploading to S3.
    The Google Drive file must be publicly accessible (shared with "Anyone with the link").
    
    Accepts various Google Drive URL formats or just the file ID:
    - File ID only: "1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H"
    - Share link: "https://drive.google.com/file/d/1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H/view"
    - Download link: "https://drive.google.com/uc?id=1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H&export=download"
    - Open link: "https://drive.google.com/open?id=1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H"
    """
    file_id: str  # Google Drive file ID or full URL (ID will be extracted automatically)
    folder_path: str  # S3 folder path (e.g., "videos/", "images/uploads/")
    filename: Optional[str] = None  # Optional filename. If not provided, will try to get from Google Drive

class TranscribeAudioRequest(BaseModel):
    audio_url: str  # Publicly accessible URL to the audio file

class AnalyzeImagesRequest(BaseModel):
    prompt: str
    image_urls: str  # Comma-separated list of image URLs
    model: str
    reasoning_effort: Optional[str] = "none"

class OptimizeImageToS3Request(BaseModel):
    image_url: str  # Publicly accessible image URL
    filename: Optional[str] = None  # Optional output filename
    output_folder: Optional[str] = "optimized-images"  # S3 folder prefix for uploads
    max_size_mb: Optional[float] = 20.0  # Max allowed image size in MB

class MarkdownToDocxRequest(BaseModel):
    text: str  # Markdown text to convert
    filename: Optional[str] = None  # Optional output filename (without extension)

class ListSharepointSubfoldersRequest(BaseModel):
    """
    Request model for listing immediate subfolders of a SharePoint folder
    via Graph using app-only auth.
    """
    sharepoint_url: str  # SharePoint URL of the parent folder
    recursive: Optional[bool] = False  # If true, list folders at all depths
    include_files: Optional[bool] = False  # If true, include files alongside folders

class CreateSharepointFolderRequest(BaseModel):
    """
    Request model for creating a folder inside a SharePoint folder via Graph
    using app-only auth.
    """
    sharepoint_url: str  # SharePoint URL of the parent folder
    folder_name: str  # Name of the new folder to create
    conflict_behavior: Optional[str] = "rename"  # "rename", "replace", or "fail"

class UploadToSharepointRequest(BaseModel):
    """
    Request model for uploading a publicly accessible file to a SharePoint folder
    via Microsoft Graph. Authentication uses an Azure AD app registration
    (client credentials flow); credentials are read from env vars
    SHAREPOINT_TENANT_ID, SHAREPOINT_CLIENT_ID, SHAREPOINT_CLIENT_SECRET.
    """
    file_url: str  # Publicly accessible URL of the file to upload
    sharepoint_url: str  # SharePoint folder URL (browser-copied URL works)
    filename: Optional[str] = None  # Optional override for the uploaded filename

app = FastAPI(title="Logic Provider Functions", version="1.0.0")

# AWS S3 configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "us-east-1")
S3_BUCKET_NAME = os.getenv("S3_BUCKET_NAME")

# Assembly AI configuration
ASSEMBLY_AI_API_KEY = os.getenv("ASSEMBLY_AI_API_KEY")

# SharePoint / Microsoft Graph (client credentials) configuration
SHAREPOINT_TENANT_ID = os.getenv("SHAREPOINT_TENANT_ID")
SHAREPOINT_CLIENT_ID = os.getenv("SHAREPOINT_CLIENT_ID")
SHAREPOINT_CLIENT_SECRET = os.getenv("SHAREPOINT_CLIENT_SECRET")
_graph_token_cache = {"token": None, "expires_at": 0.0}


def _get_graph_app_token():
    """Return a cached app-only Microsoft Graph access token."""
    now = time.time()
    if (
        _graph_token_cache["token"]
        and now < _graph_token_cache["expires_at"] - 60
    ):
        return _graph_token_cache["token"]
    if not all([SHAREPOINT_TENANT_ID, SHAREPOINT_CLIENT_ID, SHAREPOINT_CLIENT_SECRET]):
        raise HTTPException(
            status_code=500,
            detail=(
                "SharePoint app credentials not configured. Set "
                "SHAREPOINT_TENANT_ID, SHAREPOINT_CLIENT_ID, SHAREPOINT_CLIENT_SECRET."
            ),
        )
    resp = requests.post(
        f"https://login.microsoftonline.com/{SHAREPOINT_TENANT_ID}/oauth2/v2.0/token",
        data={
            "grant_type": "client_credentials",
            "client_id": SHAREPOINT_CLIENT_ID,
            "client_secret": SHAREPOINT_CLIENT_SECRET,
            "scope": "https://graph.microsoft.com/.default",
        },
        timeout=30,
    )
    if resp.status_code != 200:
        raise HTTPException(
            status_code=502,
            detail=f"Failed to obtain Graph token: {resp.status_code} {resp.text}",
        )
    j = resp.json()
    _graph_token_cache["token"] = j["access_token"]
    _graph_token_cache["expires_at"] = now + float(j.get("expires_in", 3600))
    return _graph_token_cache["token"]


def _resolve_sharepoint_folder(sharepoint_url: str):
    """
    Resolve a SharePoint folder URL to (drive_id, folder_id) via Microsoft
    Graph using app-only auth. Walks the path segment-by-segment under each
    candidate drive in the site so library/folder name guessing is robust.

    Raises HTTPException on failure.
    """
    GRAPH = "https://graph.microsoft.com/v1.0"
    auth_headers = {"Authorization": f"Bearer {_get_graph_app_token()}"}

    parsed = urlparse(sharepoint_url)
    hostname = parsed.netloc
    path = unquote(parsed.path).lstrip("/")
    path = re.sub(r"^:[a-z]:/[a-z]/", "", path)
    parts = [p for p in path.split("/") if p]
    if len(parts) < 2 or parts[0] not in ("sites", "teams"):
        raise HTTPException(
            status_code=400,
            detail=(
                f"Could not parse SharePoint URL. Expected a path containing "
                f"'sites/<name>' or 'teams/<name>'. Got: {sharepoint_url}"
            ),
        )
    site_segment = f"{parts[0]}/{parts[1]}"
    remaining = parts[2:]

    site_resp = requests.get(
        f"{GRAPH}/sites/{hostname}:/{site_segment}",
        headers=auth_headers,
        timeout=30,
    )
    if site_resp.status_code != 200:
        raise HTTPException(
            status_code=502,
            detail=f"Failed to resolve SharePoint site: {site_resp.text}",
        )
    site_id = site_resp.json().get("id")

    library_name = remaining[0] if remaining else None
    folder_path_parts = remaining[1:] if remaining else []

    drives_resp = requests.get(
        f"{GRAPH}/sites/{site_id}/drives", headers=auth_headers, timeout=30
    )
    candidate_drives = []
    seen = set()
    if drives_resp.status_code == 200:
        for d in drives_resp.json().get("value", []):
            if d.get("id") and d["id"] not in seen:
                candidate_drives.append(d)
                seen.add(d["id"])
    default_resp = requests.get(
        f"{GRAPH}/sites/{site_id}/drive", headers=auth_headers, timeout=30
    )
    if default_resp.status_code == 200:
        d = default_resp.json()
        if d.get("id") and d["id"] not in seen:
            candidate_drives.append(d)
            seen.add(d["id"])
    if not candidate_drives:
        raise HTTPException(
            status_code=502, detail="No SharePoint drives accessible for this site"
        )
    if library_name:
        def _score(d):
            name = d.get("name", "")
            web = d.get("webUrl", "").rstrip("/")
            if name == library_name:
                return 0
            if web.endswith("/" + library_name.replace(" ", "%20")):
                return 1
            return 2
        candidate_drives.sort(key=_score)

    def _walk(drive, segments):
        drv_id = drive.get("id")
        root_resp_local = requests.get(
            f"{GRAPH}/drives/{drv_id}/root", headers=auth_headers, timeout=30
        )
        if root_resp_local.status_code != 200:
            return None, f"root fetch failed: {root_resp_local.text}"
        cursor = root_resp_local.json().get("id")
        for segment in segments:
            children_url = (
                f"{GRAPH}/drives/{drv_id}/items/{cursor}/children"
                f"?$top=200&$select=id,name,folder"
            )
            found_child = None
            seen_names = []
            while children_url:
                ch_resp = requests.get(children_url, headers=auth_headers, timeout=30)
                if ch_resp.status_code != 200:
                    return None, f"children fetch failed: {ch_resp.text}"
                ch_json = ch_resp.json()
                for child in ch_json.get("value", []):
                    seen_names.append(child.get("name", ""))
                    if child.get("name", "").lower() == segment.lower():
                        found_child = child
                        break
                if found_child:
                    break
                children_url = ch_json.get("@odata.nextLink")
            if not found_child:
                return None, (
                    f"segment '{segment}' not found; available: {seen_names[:20]}"
                )
            cursor = found_child["id"]
        return cursor, None

    attempt_segments = [folder_path_parts]
    if folder_path_parts != remaining:
        attempt_segments.append(remaining)

    diagnostics = []
    for d in candidate_drives:
        for segs in attempt_segments:
            fid, err = _walk(d, segs)
            if fid:
                return d.get("id"), fid
            diagnostics.append(
                f"drive '{d.get('name')}' ({d.get('id')}) segs={segs}: {err}"
            )

    raise HTTPException(
        status_code=502,
        detail=(
            f"Failed to resolve folder path '{'/'.join(remaining)}'. "
            f"Diagnostics: {diagnostics}"
        ),
    )

# Initialize S3 client
s3_client = None
try:
    if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY and S3_BUCKET_NAME:
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
            region_name=AWS_REGION
        )
        logger.info("S3 client initialized successfully")
    else:
        logger.warning("AWS credentials not fully configured - S3 operations will fail")
except Exception as e:
    logger.error(f"Failed to initialize S3 client: {e}")
    s3_client = None

# Initialize Assembly AI client
assembly_client = None
try:
    if ASSEMBLY_AI_API_KEY:
        aai.settings.api_key = ASSEMBLY_AI_API_KEY
        logger.info("Assembly AI client initialized successfully")
    else:
        logger.warning("Assembly AI API key not configured - transcription operations will fail")
except Exception as e:
    logger.error(f"Failed to initialize Assembly AI client: {e}")
    assembly_client = None

@app.get("/")
async def root():
    logger.info("Root endpoint accessed")
    return {
        "message": "Logic Provider Functions API", 
        "status": "running",
        "endpoints": ["/upload-html", "/generate-image", "/html-to-image", "/optimize-image-to-s3", "/process-loom-video", "/upload-linkedin-video", "/url-to-google-drive", "/eleven-labs-speech", "/heygen-avatar-iv", "/google-drive-to-s3", "/telegram-file/bot{BotToken}/{file_path}", "/transcribe-audio", "/analyze-images", "/markdown-to-docx", "/upload-to-sharepoint", "/create-sharepoint-folder", "/list-sharepoint-subfolders", "/health"],
        "aws_configured": bool(AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY and S3_BUCKET_NAME),
        "timestamp": datetime.now().isoformat()
    }

@app.post("/upload-html")
async def upload_html(request: HTMLUploadRequest):
    """
    Upload HTML content to AWS S3 bucket and return public URL
    
    Args:
        request: HTMLUploadRequest containing html_content and optional filename
    
    Returns:
        JSON response with upload status and public URL
    """
    logger.info(f"Upload HTML request received - Filename: {request.filename}, Content length: {len(request.html_content)}")
    
    try:
        # Validate required environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500, 
                detail="AWS credentials or bucket name not configured"
            )
        
        # Check if S3 client is available
        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )
        
        logger.info(f"Request details: {request}")
        
        # Extract values from request
        html_content = request.html_content
        filename = request.filename
        content_type = request.content_type
        
        # Generate filename if not provided
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"html/html_{timestamp}_{unique_id}.html"
            logger.info(f"Generated filename: {filename}")
        else:
            # Ensure filename has html/ prefix and .html extension
            if not filename.startswith("html/"):
                filename = f"html/{filename}"
            if not filename.endswith('.html'):
                filename += '.html'
            logger.info(f"Adjusted filename: {filename}")
        
        logger.info(f"Attempting to upload to S3 - Bucket: {S3_BUCKET_NAME}, Key: {filename}")
        
        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=filename,
            Body=html_content.encode('utf-8'),
            ContentType=content_type,
            ACL='public-read'  # Make the object publicly readable
        )
        
        logger.info(f"Successfully uploaded to S3: {filename}")
        
        # Generate public URL
        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{filename}"
        logger.info(f"Generated public URL: {public_url}")
        
        response_data = {
            "success": True,
            "message": "HTML content uploaded successfully",
            "filename": filename,
            "public_url": public_url,
            "bucket": S3_BUCKET_NAME,
            "region": AWS_REGION,
            "key": filename,
            "uploaded_at": datetime.now().isoformat()
        }
        
        logger.info(f"Upload successful - returning response: {response_data}")
        return JSONResponse(
            status_code=200,
            content=response_data
        )
        
    except Exception as e:
        logger.error(f"Error during upload: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload HTML content: {str(e)}"
        )

@app.post("/generate-image")
async def generate_image(
    request: ImageGenerationRequest,
    authorization: Optional[str] = Header(None)
):
    """
    Generate an image using OpenAI's DALL-E API and save it to S3
    
    Args:
        request: ImageGenerationRequest containing prompt and optional size/quality
        authorization: Bearer token for OpenAI API access
    
    Returns:
        JSON response with generation status and S3 URL
    """
    logger.info(f"Image generation request received - Prompt length: {len(request.prompt)}")
    
    try:
        # Validate authorization header
        if not authorization or not authorization.lower().startswith("bearer "):
            logger.error("Missing or invalid authorization header")
            raise HTTPException(
                status_code=401,
                detail="Authorization header with Bearer token is required"
            )
        
        # Extract API key from authorization header (case insensitive)
        api_key = authorization.split(" ", 1)[1].strip()
        
        # Validate required AWS environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )
        
        # Check if S3 client is available
        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )
        
        # Prepare OpenAI API request
        openai_url = "https://api.openai.com/v1/images/generations"
        openai_headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        openai_payload = {
            "model": "gpt-image-1",
            "prompt": request.prompt,
            "size": request.size,
            "quality": request.quality
        }
        
        logger.info(f"Making request to OpenAI API with payload: {openai_payload}")
        
        # Make request to OpenAI API
        response = requests.post(
            openai_url,
            headers=openai_headers,
            json=openai_payload,
            timeout=60  # 60 second timeout for image generation
        )
        
        if response.status_code != 200:
            logger.error(f"OpenAI API error: {response.status_code} - {response.text}")
            raise HTTPException(
                status_code=response.status_code,
                detail=f"OpenAI API error: {response.text}"
            )
        
        response_data = response.json()
        logger.info(f"OpenAI API response received successfully")
        
        # Extract base64 image data
        if not response_data.get("data") or len(response_data["data"]) == 0:
            logger.error("No image data returned from OpenAI API")
            raise HTTPException(
                status_code=500,
                detail="No image data returned from OpenAI API"
            )
        
        b64_json = response_data["data"][0]["b64_json"]
        
        # Decode base64 image
        image_data = base64.b64decode(b64_json)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        filename = f"ai-generated-images/image_{timestamp}_{unique_id}.png"
        
        logger.info(f"Attempting to upload image to S3 - Bucket: {S3_BUCKET_NAME}, Key: {filename}")
        
        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=filename,
            Body=image_data,
            ContentType="image/png",
            ACL='public-read'  # Make the object publicly readable
        )
        
        logger.info(f"Successfully uploaded image to S3: {filename}")
        
        # Generate public URL
        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{filename}"
        logger.info(f"Generated public URL: {public_url}")
        
        response_result = {
            "success": True,
            "message": "Image generated and uploaded successfully",
            "public_url": public_url,
            "prompt": request.prompt,
            "size": request.size,
            "quality": request.quality,
            "generated_at": datetime.now().isoformat(),
        }
        
        logger.info(f"Image generation successful - returning response")
        return JSONResponse(
            status_code=200,
            content=response_result
        )
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error during OpenAI API call: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to connect to OpenAI API: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error during image generation: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate image: {str(e)}"
        )

@app.post("/process-loom-video")
async def process_loom_video(request: LoomVideoRequest):
    """
    Download a Loom video, upload it to S3, transcribe it with Assembly AI, and return both the video URL and transcript.

    Args:
        request: LoomVideoRequest containing loom_url and optional filename

    Returns:
        JSON response with video URL and transcription result
    """
    logger.info(f"Loom video processing request received - URL: {request.loom_url}")

    try:
        # Validate required configurations
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )

        if not ASSEMBLY_AI_API_KEY:
            logger.error("Missing Assembly AI API key")
            raise HTTPException(
                status_code=500,
                detail="Assembly AI API key not configured"
            )

        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )

        # Generate filename if not provided
        filename = request.filename
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"videos/loom_video_{timestamp}_{unique_id}.mp4"
            logger.info(f"Generated filename: {filename}")
        else:
            # Ensure it has the videos/ prefix and .mp4 extension
            if not filename.startswith("videos/"):
                filename = f"videos/{filename}"
            if not filename.endswith('.mp4'):
                filename += '.mp4'
            logger.info(f"Using provided filename: {filename}")

        # Step 1: Download video using loom-dl
        logger.info(f"Downloading video from: {request.loom_url}")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as temp_file:
            temp_video_path = temp_file.name

        try:
            # Download video using loom-dl functions
            video_id = loomdl.extract_id(request.loom_url)
            download_url = loomdl.fetch_loom_download_url(video_id)
            loomdl.download_loom_video(download_url, temp_video_path)
            logger.info(f"Video downloaded successfully to: {temp_video_path}")

            # Step 2: Upload to S3
            logger.info(f"Uploading video to S3 - Bucket: {S3_BUCKET_NAME}, Key: {filename}")
            with open(temp_video_path, 'rb') as video_file:
                s3_client.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=filename,
                    Body=video_file,
                    ContentType="video/mp4",
                    ACL='public-read'
                )
            # Generate public URL
            video_public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{filename}"
            logger.info(f"Video uploaded to S3 successfully: {video_public_url}")

            # Step 3: Send to Assembly AI for transcription
            logger.info(f"Sending video to Assembly AI for transcription: {video_public_url}")
            transcriber = aai.Transcriber()
            transcript = transcriber.transcribe(video_public_url)

            # Step 4: Poll for transcription completion
            logger.info("Polling for transcription completion...")
            max_polls = 60  # Maximum 5 minutes (60 * 5 seconds)
            poll_count = 0

            while poll_count < max_polls:
                if transcript.status == aai.TranscriptStatus.completed:
                    logger.info("Transcription completed successfully")
                    break
                elif transcript.status == aai.TranscriptStatus.error:
                    logger.error(f"Transcription failed: {transcript.error}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Transcription failed: {transcript.error}"
                    )

                logger.info(f"Transcription status: {transcript.status}, waiting...")
                time.sleep(5)  # Wait 5 seconds before polling again
                poll_count += 1

            if poll_count >= max_polls:
                logger.error("Transcription timed out")
                raise HTTPException(
                    status_code=500,
                    detail="Transcription timed out after 5 minutes"
                )

            # Clean up temporary file
            os.unlink(temp_video_path)

            # Return response
            response_data = {
                "success": True,
                "message": "Video processed and transcribed successfully",
                "video_url": video_public_url,
                "transcript": transcript.text,
                # "transcript": {
                #     "text": transcript.text,
                #     "confidence": transcript.confidence,
                #     "duration": transcript.audio_duration,
                #     "words": [{"text": word.text, "start": word.start, "end": word.end, "confidence": word.confidence}
                #              for word in transcript.words] if transcript.words else []
                # },
                "filename": filename,
                "processed_at": datetime.now().isoformat()
            }

            logger.info("Loom video processing completed successfully")
            return JSONResponse(status_code=200, content=response_data)

        except Exception as e:
            # Clean up temporary file if it exists
            if os.path.exists(temp_video_path):
                os.unlink(temp_video_path)
            raise

    except Exception as e:
        logger.error(f"Error during loom video processing: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process loom video: {str(e)}"
        )

@app.post("/html-to-image")
async def html_to_image(request: HTMLToImageRequest):
    """
    Convert HTML content to a PNG image optimized for Instagram posting.
    Uses Playwright to render HTML in a headless browser and capture as PNG.
    
    Instagram Recommended Dimensions:
    - Square Post: 1080x1080px (default)
    - Portrait Post: 1080x1350px
    - Landscape Post: 1080x566px
    - Story: 1080x1920px
    
    Args:
        request: HTMLToImageRequest containing html_content and optional dimensions/filename
    
    Returns:
        JSON response with the public S3 URL of the generated PNG image
    """
    logger.info(f"HTML-to-Image request received - HTML length: {len(request.html_content)}, Dimensions: {request.width}x{request.height}")
    
    try:
        # Validate required AWS environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )
        
        # Check if S3 client is available
        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )
        
        # Validate dimensions for Instagram
        if request.width < 320 or request.width > 1080:
            logger.warning(f"Width {request.width}px is outside Instagram's recommended range (320-1080px)")
        
        if request.height < 566 or request.height > 1920:
            logger.warning(f"Height {request.height}px is outside Instagram's recommended range (566-1920px)")
        
        # Import Playwright
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            logger.error("Playwright not installed")
            raise HTTPException(
                status_code=500,
                detail="Playwright library not installed. Please run: pip install playwright && playwright install chromium"
            )
        
        # Create temporary file for the screenshot
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
            temp_image_path = temp_file.name
        
        try:
            logger.info("Launching Playwright browser...")
            
            # Launch Playwright and render HTML
            with sync_playwright() as p:
                # Launch browser - use chromium for best compatibility
                browser = p.chromium.launch(
                    headless=True,
                    args=[
                        '--no-sandbox',
                        '--disable-setuid-sandbox',
                        '--disable-dev-shm-usage',
                        '--disable-accelerated-2d-canvas',
                        '--no-first-run',
                        '--no-zygote',
                        '--disable-gpu'
                    ]
                )
                
                logger.info("Browser launched successfully")
                
                # Create a new page with specified viewport
                page = browser.new_page(
                    viewport={'width': request.width, 'height': request.height},
                    device_scale_factor=2  # Retina display for higher quality
                )
                
                logger.info(f"Page created with viewport: {request.width}x{request.height}")
                
                # Set the HTML content
                page.set_content(request.html_content, wait_until='networkidle')
                
                logger.info("HTML content loaded successfully")
                
                # Take screenshot
                page.screenshot(
                    path=temp_image_path,
                    type='png',
                    full_page=False,  # Only capture viewport
                    omit_background=False  # Include background
                )
                
                logger.info(f"Screenshot captured: {temp_image_path}")
                
                # Close browser
                browser.close()
                logger.info("Browser closed")
            
            # Generate filename if not provided
            filename = request.filename
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                unique_id = str(uuid.uuid4())[:8]
                filename = f"instagram-images/ig_post_{timestamp}_{unique_id}.png"
            else:
                # Ensure filename has instagram-images/ prefix and .png extension
                if not filename.startswith("instagram-images/"):
                    filename = f"instagram-images/{filename}"
                if not filename.endswith('.png'):
                    filename += '.png'
            
            logger.info(f"Uploading image to S3 - Bucket: {S3_BUCKET_NAME}, Key: {filename}")
            
            # Upload to S3
            with open(temp_image_path, 'rb') as image_file:
                s3_client.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=filename,
                    Body=image_file,
                    ContentType="image/png",
                    ACL='public-read',
                    Metadata={
                        'width': str(request.width),
                        'height': str(request.height),
                        'generated-by': 'html-to-image-api'
                    }
                )
            
            logger.info(f"Image uploaded to S3 successfully: {filename}")
            
            # Generate public URL
            public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{filename}"
            logger.info(f"Generated public URL: {public_url}")
            
            # Clean up temporary file
            os.unlink(temp_image_path)
            
            # Return response
            response_data = {
                "success": True,
                "message": "HTML converted to image successfully",
                "public_url": public_url,
                "filename": filename,
                "dimensions": {
                    "width": request.width,
                    "height": request.height
                },
                "instagram_optimized": True,
                "format": "PNG",
                "generated_at": datetime.now().isoformat()
            }
            
            logger.info("HTML-to-Image conversion completed successfully")
            return JSONResponse(status_code=200, content=response_data)
            
        except Exception as e:
            # Clean up temporary file if it exists
            if os.path.exists(temp_image_path):
                os.unlink(temp_image_path)
            raise
            
    except Exception as e:
        logger.error(f"Error during HTML-to-Image conversion: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to convert HTML to image: {str(e)}"
        )

@app.post("/optimize-image-to-s3")
async def optimize_image_to_s3(request: OptimizeImageToS3Request):
    """
    Download a public image URL and ensure uploaded size is <= max_size_mb.
    If the source image is already under the limit, upload as-is.
    If it exceeds the limit, compress/resize and upload the optimized image.
    """
    logger.info(f"Optimize image request received - URL: {request.image_url}")

    try:
        # Validate required AWS environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )

        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )

        max_size_mb = request.max_size_mb if request.max_size_mb and request.max_size_mb > 0 else 20.0
        max_size_bytes = int(max_size_mb * 1024 * 1024)

        # Step 1: Download image from public URL
        logger.info("Downloading image from URL...")
        response = requests.get(
            request.image_url,
            timeout=120,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            }
        )
        response.raise_for_status()

        original_bytes = response.content
        original_size_bytes = len(original_bytes)
        if original_size_bytes == 0:
            raise HTTPException(status_code=400, detail="Downloaded image is empty")

        source_content_type = response.headers.get("Content-Type", "application/octet-stream")
        if ";" in source_content_type:
            source_content_type = source_content_type.split(";")[0].strip()

        logger.info(
            f"Downloaded image size: {original_size_bytes} bytes ({original_size_bytes / (1024 * 1024):.2f} MB), Content-Type: {source_content_type}"
        )

        # Step 2: Resolve output filename
        filename = request.filename
        if not filename:
            content_disposition = response.headers.get("Content-Disposition", "")
            if "filename=" in content_disposition:
                match = re.search(r'filename\*?=["\']?(?:UTF-8\'\')?([^"\';\s]+)["\']?', content_disposition)
                if match:
                    filename = unquote(match.group(1))

        if not filename:
            parsed = urlparse(request.image_url)
            path_filename = unquote(parsed.path.split("/")[-1]) if parsed.path else ""
            if path_filename:
                filename = path_filename

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            extension_map = {
                "image/jpeg": ".jpg",
                "image/png": ".png",
                "image/webp": ".webp",
                "image/gif": ".gif",
                "image/svg+xml": ".svg",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
            }
            extension = extension_map.get(source_content_type, ".jpg")
            filename = f"image_{timestamp}_{unique_id}{extension}"

        # Basic filename safety: avoid writing nested key paths via filename input.
        filename = filename.replace("\\", "/").split("/")[-1]

        # Normalize folder path
        output_folder = (request.output_folder or "optimized-images").strip("/")
        if not output_folder:
            output_folder = "optimized-images"

        is_svg = source_content_type == "image/svg+xml" or filename.lower().endswith(".svg")
        upload_bytes = original_bytes
        upload_content_type = source_content_type
        was_compressed = False

        # Step 3: If over max size, compress/resize for raster images.
        if original_size_bytes > max_size_bytes:
            if is_svg:
                raise HTTPException(
                    status_code=400,
                    detail="SVG image is larger than the max size and cannot be resized by this endpoint"
                )

            try:
                image = Image.open(BytesIO(original_bytes))
                image.load()
            except Exception as image_error:
                raise HTTPException(
                    status_code=400,
                    detail=f"Downloaded file is not a valid processable image: {str(image_error)}"
                )

            # Convert to RGB for JPEG encoding (flatten transparency onto white if needed).
            if image.mode in ("RGBA", "LA", "P"):
                rgba_image = image.convert("RGBA")
                background = Image.new("RGBA", rgba_image.size, (255, 255, 255, 255))
                image = Image.alpha_composite(background, rgba_image).convert("RGB")
            else:
                image = image.convert("RGB")

            resample_filter = Image.Resampling.LANCZOS if hasattr(Image, "Resampling") else Image.LANCZOS
            quality = 92
            resize_ratio = 1.0
            best_result = None

            for _ in range(14):
                working_image = image
                if resize_ratio < 1.0:
                    new_width = max(1, int(image.width * resize_ratio))
                    new_height = max(1, int(image.height * resize_ratio))
                    working_image = image.resize((new_width, new_height), resample_filter)

                output_buffer = BytesIO()
                working_image.save(
                    output_buffer,
                    format="JPEG",
                    quality=quality,
                    optimize=True,
                    progressive=True
                )
                candidate_bytes = output_buffer.getvalue()
                best_result = candidate_bytes

                if len(candidate_bytes) <= max_size_bytes:
                    break

                if quality > 45:
                    quality -= 10
                else:
                    resize_ratio *= 0.85

            if not best_result or len(best_result) > max_size_bytes:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unable to reduce image below {max_size_mb} MB"
                )

            base_name, _ = os.path.splitext(filename)
            filename = f"{base_name}_optimized.jpg"
            upload_bytes = best_result
            upload_content_type = "image/jpeg"
            was_compressed = True
        else:
            # Validate non-SVG images are processable image files.
            if not is_svg:
                try:
                    image = Image.open(BytesIO(original_bytes))
                    image.load()
                except Exception as image_error:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Downloaded file is not a valid image: {str(image_error)}"
                    )

        # If no extension is present, infer one from upload content type.
        if "." not in filename:
            extension_map = {
                "image/jpeg": ".jpg",
                "image/png": ".png",
                "image/webp": ".webp",
                "image/gif": ".gif",
                "image/svg+xml": ".svg",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
            }
            filename += extension_map.get(upload_content_type, ".jpg")

        s3_key = f"{output_folder}/{filename}"
        logger.info(f"Uploading optimized image to S3 - Bucket: {S3_BUCKET_NAME}, Key: {s3_key}")

        # Step 4: Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=s3_key,
            Body=upload_bytes,
            ContentType=upload_content_type,
            ContentDisposition="inline",
            ACL='public-read'
        )

        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
        final_size_bytes = len(upload_bytes)
        logger.info(
            f"Image upload complete. Final size: {final_size_bytes} bytes ({final_size_bytes / (1024 * 1024):.2f} MB)"
        )

        response_data = {
            "success": True,
            "message": "Image uploaded successfully (optimized only when required)",
            "source_url": request.image_url,
            "public_url": public_url,
            "s3_key": s3_key,
            "filename": filename,
            "folder_path": output_folder,
            "max_size_mb": max_size_mb,
            "original_size_bytes": original_size_bytes,
            "final_size_bytes": final_size_bytes,
            "was_compressed": was_compressed,
            "content_type": upload_content_type,
            "uploaded_at": datetime.now().isoformat()
        }
        return JSONResponse(status_code=200, content=response_data)

    except requests.exceptions.RequestException as e:
        logger.error(f"Failed to download image from URL: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Failed to download image from URL: {str(e)}"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during image optimization: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to optimize and upload image: {str(e)}"
        )

@app.post("/upload-linkedin-video")
async def upload_linkedin_video(
    request: LinkedInVideoRequest,
    authorization: str = Header(...)
):
    """
    Upload a video to LinkedIn and create a UGC post with caption using LinkedIn v2 API.

    Video Requirements:
    - Format: MP4
    - Duration: 3 seconds to 30 minutes
    - Size: 75 KB to 500 MB

    The LinkedIn access token must have the following permissions:
    - r_liteprofile: Read basic profile information
    - w_member_social: Write posts and content

    Uses LinkedIn v2 API endpoints in a 3-step process:
    - POST /v2/assets?action=registerUpload (initialize upload)
    - PUT to upload URL (upload video file)
    - POST /v2/ugcPosts (create UGC post)

    Args:
        request: LinkedInVideoRequest containing video_url, caption, and optional visibility/filename
        authorization: Bearer token for LinkedIn API access (should be "Bearer <token>")

    Returns:
        JSON response with upload status and post details
    """
    logger.info(f"LinkedIn video upload request received - Video URL: {request.video_url}")

    try:
        # Validate authorization header
        if not authorization or not authorization.lower().startswith("bearer "):
            logger.error("Missing or invalid authorization header")
            raise HTTPException(
                status_code=401,
                detail="Authorization header with Bearer token is required"
            )

        # Extract access token
        access_token = authorization.split(" ", 1)[1].strip()

        # Validate visibility setting
        valid_visibilities = ["PUBLIC", "CONNECTIONS", "LOGGED_IN"]
        if request.visibility.upper() not in valid_visibilities:
            logger.error(f"Invalid visibility setting: {request.visibility}")
            raise HTTPException(
                status_code=400,
                detail=f"Invalid visibility. Must be one of: {', '.join(valid_visibilities)}"
            )

        # Step 1: Download the video from the provided URL
        logger.info(f"Downloading video from: {request.video_url}")
        video_response = requests.get(request.video_url, timeout=300)  # 5 minute timeout for large videos
        video_response.raise_for_status()

        video_content = video_response.content
        content_length = len(video_content)

        # Validate video specifications per LinkedIn requirements
        min_size_kb = 75 * 1024  # 75 KB
        max_size_mb = 500 * 1024 * 1024  # 500 MB

        if content_length < min_size_kb:
            logger.error(f"Video too small: {content_length} bytes. Minimum: {min_size_kb} bytes (75 KB)")
            raise HTTPException(
                status_code=400,
                detail="Video file is too small. LinkedIn requires minimum 75 KB"
            )

        if content_length > max_size_mb:
            logger.error(f"Video too large: {content_length} bytes. Maximum: {max_size_mb} bytes (500 MB)")
            raise HTTPException(
                status_code=400,
                detail="Video file is too large. LinkedIn allows maximum 500 MB"
            )

        logger.info(f"Video downloaded successfully. Size: {content_length} bytes ({content_length / (1024*1024):.2f} MB)")

        # Generate filename if not provided
        filename = request.filename
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"linkedin_video_{timestamp}_{unique_id}.mp4"
        elif not filename.endswith('.mp4'):
            filename += '.mp4'

        # Step 2: Get user's member ID from LinkedIn
        logger.info("Getting user member ID from LinkedIn API")

        linkedin_headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
            "X-Restli-Protocol-Version": "2.0.0",
            "LinkedIn-Version": "202409"  # Required for LinkedIn REST API
        }

        # Get the authenticated user's info to get their member ID
        profile_url = "https://api.linkedin.com/v2/userinfo"
        profile_response = requests.get(
            profile_url,
            headers=linkedin_headers,
            timeout=30
        )

        if profile_response.status_code != 200:
            logger.error(f"Failed to get LinkedIn userinfo: {profile_response.status_code} - {profile_response.text}")
            raise HTTPException(
                status_code=profile_response.status_code,
                detail=f"Failed to get LinkedIn user information: {profile_response.text}"
            )

        profile_data = profile_response.json()
        member_id = profile_data.get("sub")

        if not member_id:
            logger.error(f"Could not extract member ID from userinfo response: {profile_data}")
            raise HTTPException(
                status_code=500,
                detail="Could not extract member ID from LinkedIn userinfo"
            )

        logger.info(f"Retrieved member ID: {member_id}")

        # Step 3: Initialize the Upload with LinkedIn v2 API
        logger.info("Initializing video upload with LinkedIn v2 API")

        register_payload = {
            "registerUploadRequest": {
                "recipes": [
                    "urn:li:digitalmediaRecipe:feedshare-video"
                ],
                "owner": f"urn:li:person:{member_id}",
                "serviceRelationships": [
                    {
                        "relationshipType": "OWNER",
                        "identifier": "urn:li:userGeneratedContent"
                    }
                ]
            }
        }

        register_url = "https://api.linkedin.com/v2/assets?action=registerUpload"
        register_response = requests.post(
            register_url,
            headers=linkedin_headers,
            json=register_payload,
            timeout=30
        )

        if register_response.status_code != 200:
            logger.error(f"LinkedIn video registration failed: {register_response.status_code} - {register_response.text}")
            raise HTTPException(
                status_code=register_response.status_code,
                detail=f"Failed to register video upload: {register_response.text}"
            )

        register_data = register_response.json()

        # Extract asset URN and upload URL from v2 API response
        asset_urn = register_data.get("value", {}).get("asset")

        upload_mechanism = register_data.get("value", {}).get("uploadMechanism", {})
        upload_url = upload_mechanism.get("com.linkedin.digitalmedia.uploading.MediaUploadHttpRequest", {}).get("uploadUrl")

        if not asset_urn or not upload_url:
            logger.error(f"Invalid response from LinkedIn v2 API registration: {register_data}")
            raise HTTPException(
                status_code=500,
                detail="Invalid response from LinkedIn v2 API during registration"
            )

        logger.info(f"Video registered successfully. Asset URN: {asset_urn}")

        # Step 4: Upload the Video File
        logger.info(f"Uploading video to LinkedIn upload URL: {upload_url}")

        # Note: Do NOT include Authorization header in this PUT request
        upload_headers = {
            "Content-Type": "video/mp4"
        }

        upload_response = requests.put(
            upload_url,
            headers=upload_headers,
            data=video_content,
            timeout=600  # 10 minute timeout for upload
        )

        if upload_response.status_code not in [200, 201]:
            logger.error(f"LinkedIn video upload failed: {upload_response.status_code} - {upload_response.text}")
            raise HTTPException(
                status_code=upload_response.status_code,
                detail=f"Failed to upload video: {upload_response.text}"
            )

        logger.info("Video uploaded successfully to LinkedIn")

        # Step 5: Create the Post (UGC Post)
        logger.info("Creating LinkedIn UGC post with video")

        post_payload = {
            "author": f"urn:li:person:{member_id}",
            "lifecycleState": "PUBLISHED",
            "specificContent": {
                "com.linkedin.ugc.ShareContent": {
                    "shareCommentary": {
                        "text": request.caption
                    },
                    "shareMediaCategory": "VIDEO",
                    "media": [
                        {
                            "status": "READY",
                            "description": {
                                "text": "Video uploaded via API"
                            },
                            "media": asset_urn
                        }
                    ]
                }
            },
            "visibility": {
                "com.linkedin.ugc.MemberNetworkVisibility": request.visibility.upper()
            }
        }

        post_url = "https://api.linkedin.com/v2/ugcPosts"
        post_response = requests.post(
            post_url,
            headers=linkedin_headers,
            json=post_payload,
            timeout=30
        )

        if post_response.status_code not in [200, 201]:
            logger.error(f"LinkedIn UGC post creation failed: {post_response.status_code} - {post_response.text}")
            raise HTTPException(
                status_code=post_response.status_code,
                detail=f"Failed to create UGC post: {post_response.text}"
            )

        post_data = post_response.json()
        post_urn = post_data.get("id")

        if not post_urn:
            logger.warning(f"No post ID in UGC Posts API response: {post_data}")
            post_urn = post_data.get("urn") or f"urn:li:share:{post_data.get('id', 'unknown')}"

        logger.info(f"LinkedIn UGC post created successfully. Post URN: {post_urn}")

        # Return success response
        response_data = {
            "success": True,
            "message": "Video uploaded and posted to LinkedIn successfully",
            "asset_urn": asset_urn,
            "post_urn": post_urn,
            "caption": request.caption,
            "visibility": request.visibility,
            "filename": filename,
            "uploaded_at": datetime.now().isoformat()
        }

        logger.info("LinkedIn video upload completed successfully")
        return JSONResponse(status_code=200, content=response_data)

    except requests.exceptions.RequestException as e:
        logger.error(f"Request error during LinkedIn API call: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to connect to LinkedIn API: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error during LinkedIn video upload: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload video to LinkedIn: {str(e)}"
        )

@app.post("/url-to-google-drive")
async def url_to_google_drive(
    request: UrlToGoogleDriveRequest,
    authorization: str = Header(...)
):
    """
    Download a file from a publicly accessible URL and upload it to Google Drive.
    
    This endpoint acts as a middleman that:
    1. Downloads the file from the provided URL
    2. Uploads it to the specified Google Drive folder using the provided credentials
    
    Args:
        request: UrlToGoogleDriveRequest containing file_url, folder_id, and optional filename
        authorization: Bearer token for Google Drive API access (should be "Bearer <access_token>")
    
    Returns:
        JSON response with upload status and Google Drive file details
    """
    logger.info(f"URL to Google Drive request received - URL: {request.file_url}, Folder ID: {request.folder_id}")
    
    try:
        # Validate authorization header
        if not authorization or not authorization.lower().startswith("bearer "):
            logger.error("Missing or invalid authorization header")
            raise HTTPException(
                status_code=401,
                detail="Authorization header with Bearer token is required. Format: 'Bearer <google_drive_access_token>'"
            )
        
        # Extract access token
        access_token = authorization.split(" ", 1)[1].strip()
        
        # Step 1: Download the file from the provided URL
        logger.info(f"Downloading file from: {request.file_url}")
        
        try:
            download_response = requests.get(
                request.file_url,
                timeout=300,  # 5 minute timeout for large files
                stream=True
            )
            download_response.raise_for_status()
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to download file from URL: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to download file from URL: {str(e)}"
            )
        
        file_content = download_response.content
        content_length = len(file_content)
        
        # Try to determine content type from response headers
        content_type = download_response.headers.get('Content-Type', 'application/octet-stream')
        # Clean up content type (remove charset and other parameters)
        if ';' in content_type:
            content_type = content_type.split(';')[0].strip()
        
        logger.info(f"File downloaded successfully. Size: {content_length} bytes ({content_length / (1024*1024):.2f} MB), Content-Type: {content_type}")
        
        # Step 2: Determine filename
        filename = request.filename
        if not filename:
            # Try to extract filename from URL or Content-Disposition header
            content_disposition = download_response.headers.get('Content-Disposition', '')
            if 'filename=' in content_disposition:
                match = re.search(r'filename[*]?=["\']?([^"\';\s]+)["\']?', content_disposition)
                if match:
                    filename = match.group(1)
            
            if not filename:
                # Extract from URL path
                parsed_url = urlparse(request.file_url)
                path_filename = parsed_url.path.split('/')[-1]
                if path_filename:
                    filename = unquote(path_filename)
                else:
                    # Generate a default filename
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    unique_id = str(uuid.uuid4())[:8]
                    # Try to determine extension from content type
                    extension_map = {
                        'image/jpeg': '.jpg',
                        'image/png': '.png',
                        'image/gif': '.gif',
                        'image/webp': '.webp',
                        'video/mp4': '.mp4',
                        'video/webm': '.webm',
                        'application/pdf': '.pdf',
                        'text/plain': '.txt',
                        'text/html': '.html',
                        'application/json': '.json',
                    }
                    extension = extension_map.get(content_type, '')
                    filename = f"file_{timestamp}_{unique_id}{extension}"
        
        logger.info(f"Using filename: {filename}")
        
        # Step 3: Upload to Google Drive using the REST API
        logger.info(f"Uploading file to Google Drive folder: {request.folder_id}")
        
        # Google Drive API v3 - multipart upload
        # First, create metadata
        metadata = {
            "name": filename,
            "parents": [request.folder_id]
        }
        
        # Use multipart upload for simplicity
        import io
        from email.mime.multipart import MIMEMultipart
        from email.mime.base import MIMEBase
        from email.mime.application import MIMEApplication
        
        # Create the multipart request manually
        boundary = f"----WebKitFormBoundary{uuid.uuid4().hex[:16]}"
        
        # Build the multipart body
        body_parts = []
        
        # Part 1: Metadata (JSON)
        body_parts.append(f'--{boundary}'.encode())
        body_parts.append(b'Content-Type: application/json; charset=UTF-8')
        body_parts.append(b'')
        body_parts.append(json.dumps(metadata).encode())
        
        # Part 2: File content
        body_parts.append(f'--{boundary}'.encode())
        body_parts.append(f'Content-Type: {content_type}'.encode())
        body_parts.append(b'')
        body_parts.append(file_content)
        
        # Closing boundary
        body_parts.append(f'--{boundary}--'.encode())
        
        # Join with CRLF
        body = b'\r\n'.join(body_parts)
        
        # Make the upload request
        upload_url = "https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart"
        
        upload_headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": f"multipart/related; boundary={boundary}"
        }
        
        upload_response = requests.post(
            upload_url,
            headers=upload_headers,
            data=body,
            timeout=600  # 10 minute timeout for large uploads
        )
        
        if upload_response.status_code not in [200, 201]:
            logger.error(f"Google Drive upload failed: {upload_response.status_code} - {upload_response.text}")
            raise HTTPException(
                status_code=upload_response.status_code,
                detail=f"Failed to upload to Google Drive: {upload_response.text}"
            )
        
        drive_response = upload_response.json()
        file_id = drive_response.get("id")
        file_name = drive_response.get("name")
        
        logger.info(f"File uploaded to Google Drive successfully. File ID: {file_id}")
        
        # Generate Google Drive file URL
        drive_file_url = f"https://drive.google.com/file/d/{file_id}/view"
        
        # Return success response
        response_data = {
            "success": True,
            "message": "File downloaded and uploaded to Google Drive successfully",
            "file_id": file_id,
            "file_name": file_name,
            "drive_url": drive_file_url,
            "folder_id": request.folder_id,
            "source_url": request.file_url,
            "file_size_bytes": content_length,
            "content_type": content_type,
            "uploaded_at": datetime.now().isoformat()
        }
        
        logger.info("URL to Google Drive transfer completed successfully")
        return JSONResponse(status_code=200, content=response_data)
    
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error during Google Drive API call: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to connect to Google Drive API: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error during URL to Google Drive transfer: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to transfer file to Google Drive: {str(e)}"
        )

@app.post("/eleven-labs-speech")
async def eleven_labs_speech(
    request: ElevenLabsSpeechRequest,
    xi_api_key: str = Header(..., alias="xi-api-key")
):
    """
    Generate speech using Eleven Labs Text-to-Speech API and upload to S3.
    
    This endpoint:
    1. Takes the xi-api-key from request headers
    2. Calls Eleven Labs API to generate speech from text
    3. Uploads the returned audio file to S3 under the audio/ folder
    4. Returns the public URL of the audio file
    
    Args:
        request: ElevenLabsSpeechRequest containing voice_id, text, optional model_id and filename
        xi_api_key: Eleven Labs API key passed via xi-api-key header
    
    Returns:
        JSON response with the public S3 URL of the generated audio file
    """
    logger.info(f"Eleven Labs speech request received - Voice ID: {request.voice_id}, Text length: {len(request.text)}")
    
    try:
        # Validate required AWS environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )
        
        # Check if S3 client is available
        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )
        
        # Validate text is not empty
        if not request.text.strip():
            logger.error("Empty text provided")
            raise HTTPException(
                status_code=400,
                detail="Text cannot be empty"
            )
        
        # Prepare Eleven Labs API request
        eleven_labs_url = f"https://api.elevenlabs.io/v1/text-to-speech/{request.voice_id}"
        eleven_labs_headers = {
            "xi-api-key": xi_api_key,
            "Content-Type": "application/json",
            "Accept": "audio/mpeg"
        }
        
        # Build voice_settings object with provided values or defaults
        voice_settings = {}
        
        # Stability: 0.0-1.0, default 0.5
        voice_settings["stability"] = request.stability if request.stability is not None else 0.5
        
        # Similarity boost: 0.0-1.0, default 0.75
        voice_settings["similarity_boost"] = request.similarity_boost if request.similarity_boost is not None else 0.75
        
        # Style: 0.0-1.0, default 0 (only add if provided since it increases latency)
        if request.style is not None:
            voice_settings["style"] = request.style
        
        # Use speaker boost: default true (only add if explicitly set)
        if request.use_speaker_boost is not None:
            voice_settings["use_speaker_boost"] = request.use_speaker_boost
        
        # Speed: default 1.0 (only add if provided)
        if request.speed is not None:
            voice_settings["speed"] = request.speed
        
        eleven_labs_payload = {
            "text": request.text,
            "model_id": request.model_id,
            "voice_settings": voice_settings
        }
        
        logger.info(f"Making request to Eleven Labs API for voice: {request.voice_id}, voice_settings: {voice_settings}")
        
        # Make request to Eleven Labs API
        response = requests.post(
            eleven_labs_url,
            headers=eleven_labs_headers,
            json=eleven_labs_payload,
            timeout=120  # 2 minute timeout for speech generation
        )
        
        if response.status_code != 200:
            logger.error(f"Eleven Labs API error: {response.status_code} - {response.text}")
            raise HTTPException(
                status_code=response.status_code,
                detail=f"Eleven Labs API error: {response.text}"
            )
        
        # Get the audio content
        audio_content = response.content
        content_length = len(audio_content)
        
        # Calculate audio duration using mutagen
        audio_duration_seconds = None
        try:
            audio_file = BytesIO(audio_content)
            audio = MP3(audio_file)
            audio_duration_seconds = round(audio.info.length, 2)  # Round to 2 decimal places
            logger.info(f"Audio duration: {audio_duration_seconds} seconds")
        except Exception as e:
            logger.warning(f"Could not determine audio duration: {str(e)}")
        
        logger.info(f"Audio generated successfully. Size: {content_length} bytes ({content_length / 1024:.2f} KB)")
        
        # Generate filename if not provided
        filename = request.filename
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"audio/speech_{timestamp}_{unique_id}.mp3"
        else:
            # Ensure filename has audio/ prefix and .mp3 extension
            if not filename.startswith("audio/"):
                filename = f"audio/{filename}"
            if not filename.endswith('.mp3'):
                filename += '.mp3'
        
        logger.info(f"Uploading audio to S3 - Bucket: {S3_BUCKET_NAME}, Key: {filename}")
        
        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=filename,
            Body=audio_content,
            ContentType="audio/mpeg",
            ACL='public-read'
        )
        
        logger.info(f"Audio uploaded to S3 successfully: {filename}")
        
        # Generate public URL
        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{filename}"
        logger.info(f"Generated public URL: {public_url}")
        
        # Return success response
        response_data = {
            "success": True,
            "message": "Speech generated and uploaded successfully",
            "public_url": public_url,
            "filename": filename,
            "voice_id": request.voice_id,
            "model_id": request.model_id,
            "voice_settings": voice_settings,
            "text_length": len(request.text),
            "audio_size_bytes": content_length,
            "audio_duration_seconds": audio_duration_seconds,
            "generated_at": datetime.now().isoformat()
        }
        
        logger.info("Eleven Labs speech generation completed successfully")
        return JSONResponse(status_code=200, content=response_data)
    
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error during Eleven Labs API call: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to connect to Eleven Labs API: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error during speech generation: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate speech: {str(e)}"
        )

@app.post("/heygen-avatar-iv")
async def heygen_avatar_iv(
    request: HeyGenAvatarIVRequest,
    x_api_key: str = Header(..., alias="X-Api-Key")
):
    """
    Generate an Avatar IV video using HeyGen's v2/video/generate API.
    
    This endpoint:
    1. Takes a publicly accessible image URL
    2. Uploads the image to HeyGen using their Upload Asset API
    3. Uses the uploaded image to generate an Avatar IV video
    4. Returns the response from HeyGen
    
    HeyGen API Documentation:
    - Upload Asset: https://docs.heygen.com/reference/upload-asset
    - Create a Video (V2): https://docs.heygen.com/reference/create-an-avatar-video-v2
    
    Args:
        request: HeyGenAvatarIVRequest containing:
            - image_url: Publicly accessible URL of the image (required)
            - script: Text for the avatar to speak (required)
            - voice_id: Voice ID to use (required)
            - video_title: Optional title for the video
            - video_orientation: Optional video orientation
            - fit: Optional fit mode ("cover" or "contain")
            - custom_motion_prompt: Optional custom motion prompt
            - enhance_custom_motion_prompt: Optional boolean to enhance motion prompt
            - audio_url: Optional URL of audio file to use instead of TTS
            - audio_asset_id: Optional asset ID of uploaded audio
        x_api_key: HeyGen API key passed via X-Api-Key header
    
    Returns:
        JSON response from HeyGen's Create a Video (V2) API
    """
    logger.info(f"HeyGen Avatar IV request received - Image URL: {request.image_url}, Script: {'Yes' if request.script else 'No'}, Audio URL: {'Yes' if request.audio_url else 'No'}")
    
    try:
        # Validate voice input mode:
        # Mode A (text): script + voice_id
        # Mode B (audio): audio_url xor audio_asset_id
        has_text_mode = bool(request.script and request.voice_id)
        has_audio_url = bool(request.audio_url)
        has_audio_asset_id = bool(request.audio_asset_id)
        has_audio_mode = has_audio_url or has_audio_asset_id

        if not has_text_mode and not has_audio_mode:
            raise HTTPException(
                status_code=400,
                detail="Provide either (script + voice_id) for text mode OR (audio_url or audio_asset_id) for audio mode"
            )
        if has_audio_url and has_audio_asset_id:
            raise HTTPException(
                status_code=400,
                detail="Provide only one of audio_url or audio_asset_id"
            )

        # Step 1: Download the image from the provided URL
        logger.info(f"Downloading image from URL: {request.image_url}")
        
        try:
            # Use a proper User-Agent header and handle URL encoding
            download_headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            }
            image_response = requests.get(
                request.image_url,
                headers=download_headers,
                timeout=60,  # 1 minute timeout for image download
                stream=True
            )
            image_response.raise_for_status()
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to download image from URL: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to download image from URL: {str(e)}"
            )
        
        image_content = image_response.content
        content_length = len(image_content)
        
        if content_length == 0:
            logger.error("Downloaded image is empty")
            raise HTTPException(
                status_code=400,
                detail="Downloaded image is empty. Please check the image URL."
            )
        
        # Determine content type from response headers
        content_type = image_response.headers.get('Content-Type', 'image/png')
        if ';' in content_type:
            content_type = content_type.split(';')[0].strip()
        
        # Determine file extension based on content type
        extension_map = {
            'image/jpeg': '.jpg',
            'image/jpg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'image/webp': '.webp',
        }
        extension = extension_map.get(content_type, '.png')
        
        logger.info(f"Image downloaded successfully. Size: {content_length} bytes, Content-Type: {content_type}")
        
        # Step 2: Upload the image to HeyGen as raw binary data
        logger.info("Uploading image to HeyGen...")
        
        upload_url = "https://upload.heygen.com/v1/asset"
        upload_headers = {
            "X-Api-Key": x_api_key,
            "Content-Type": content_type  # Set content type to actual image type (e.g., image/png)
        }
        
        # Upload as raw binary data (not multipart form-data)
        upload_response = requests.post(
            upload_url,
            headers=upload_headers,
            data=image_content,  # Raw binary data
            timeout=120  # 2 minute timeout for upload
        )
        
        logger.info(f"HeyGen upload response status: {upload_response.status_code}")
        logger.info(f"HeyGen upload response body: {upload_response.text}")
        
        if upload_response.status_code != 200:
            logger.error(f"HeyGen asset upload failed: {upload_response.status_code} - {upload_response.text}")
            raise HTTPException(
                status_code=upload_response.status_code,
                detail=f"Failed to upload image to HeyGen: {upload_response.text}"
            )
        
        upload_data = upload_response.json()
        logger.info(f"HeyGen upload response parsed: {upload_data}")
        
        # Extract the image key from the response
        # HeyGen returns the asset_id in data.image_key or data.asset_id or data.url
        image_key = None
        if "data" in upload_data:
            image_key = upload_data["data"].get("image_key") or upload_data["data"].get("asset_id") or upload_data["data"].get("url")
        elif "image_key" in upload_data:
            image_key = upload_data["image_key"]
        elif "asset_id" in upload_data:
            image_key = upload_data["asset_id"]
        
        if not image_key:
            logger.error(f"Could not extract image_key from HeyGen upload response: {upload_data}")
            raise HTTPException(
                status_code=500,
                detail=f"Could not extract image_key from HeyGen upload response: {upload_data}"
            )
        
        logger.info(f"Image uploaded to HeyGen successfully. Image key: {image_key}")

        # Step 3: Create a photo avatar look from uploaded image_key.
        # /v2/video/generate requires a valid talking_photo_id (avatar/look id), not raw image_key.
        logger.info("Creating HeyGen photo avatar look from uploaded image...")
        create_look_url = "https://api.heygen.com/v2/photo_avatar/avatar_group/create"
        look_name = request.video_title or f"avatar_iv_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        create_look_payload = {
            "name": look_name,
            "image_key": image_key
        }
        heygen_json_headers = {
            "X-Api-Key": x_api_key,
            "Content-Type": "application/json"
        }

        create_look_response = requests.post(
            create_look_url,
            headers=heygen_json_headers,
            json=create_look_payload,
            timeout=120
        )

        if create_look_response.status_code != 200:
            logger.error(f"HeyGen create photo avatar look failed: {create_look_response.status_code} - {create_look_response.text}")
            raise HTTPException(
                status_code=create_look_response.status_code,
                detail=f"Failed to create HeyGen photo avatar look: {create_look_response.text}"
            )

        create_look_data = create_look_response.json()
        talking_photo_id = (create_look_data.get("data") or {}).get("id")
        if not talking_photo_id:
            logger.error(f"Could not extract talking_photo_id from create look response: {create_look_data}")
            raise HTTPException(
                status_code=500,
                detail=f"Could not extract talking_photo_id from HeyGen create look response: {create_look_data}"
            )

        logger.info(f"Created photo avatar look id: {talking_photo_id}. Polling until ready...")

        # Step 4: Poll look status until completed (or timeout/failure)
        details_url = f"https://api.heygen.com/v2/photo_avatar/{talking_photo_id}"
        max_polls = 30
        poll_interval_seconds = 4
        look_status = None

        for _ in range(max_polls):
            details_response = requests.get(details_url, headers={"X-Api-Key": x_api_key}, timeout=60)
            if details_response.status_code != 200:
                logger.warning(f"HeyGen look status poll failed: {details_response.status_code} - {details_response.text}")
                time.sleep(poll_interval_seconds)
                continue

            details_data = details_response.json()
            look_status = ((details_data.get("data") or {}).get("status") or "").lower()
            if look_status == "completed":
                logger.info(f"HeyGen look is ready: {talking_photo_id}")
                break
            if look_status in {"moderation_rejected", "in_appeal", "failed"}:
                logger.error(f"HeyGen look not usable. Status: {look_status}, details: {details_data}")
                raise HTTPException(
                    status_code=400,
                    detail=f"HeyGen look status is '{look_status}'. Details: {details_data}"
                )

            time.sleep(poll_interval_seconds)

        if look_status != "completed":
            logger.error(f"Timed out waiting for HeyGen look readiness. Last status: {look_status}, look_id: {talking_photo_id}")
            raise HTTPException(
                status_code=504,
                detail=f"Timed out waiting for HeyGen talking photo to be ready. look_id: {talking_photo_id}, last_status: {look_status}"
            )

        # Step 5: Generate video using HeyGen v2 endpoint
        logger.info("Generating Avatar IV video via v2/video/generate...")
        
        generate_url = "https://api.heygen.com/v2/video/generate"
        generate_headers = heygen_json_headers

        # Build voice object for v2 schema
        if has_text_mode:
            voice_object = {
                "type": "text",
                "voice_id": request.voice_id,
                "input_text": request.script
            }
        else:
            voice_object = {"type": "audio"}
            if request.audio_url is not None:
                voice_object["audio_url"] = request.audio_url
            if request.audio_asset_id is not None:
                voice_object["audio_asset_id"] = request.audio_asset_id

        # Core v2 payload
        generate_payload = {
            "title": request.video_title or "Avatar IV Video",
            "video_inputs": [
                {
                    "character": {
                        "type": "talking_photo",
                        "talking_photo_id": talking_photo_id,
                        "talking_style": "expressive",
                        "use_avatar_iv_model": True,
                        "super_resolution": request.super_resolution if request.super_resolution is not None else True
                    },
                    "voice": voice_object
                }
            ]
        }

        # Optional output dimensions derived from old orientation field
        if request.video_orientation:
            orientation = request.video_orientation.lower()
            if orientation == "portrait":
                generate_payload["dimension"] = {"width": 1080, "height": 1920}
            elif orientation == "square":
                generate_payload["dimension"] = {"width": 1080, "height": 1080}
            elif orientation == "landscape":
                generate_payload["dimension"] = {"width": 1920, "height": 1080}
        
        logger.info(f"Calling HeyGen Avatar IV API with payload: {generate_payload}")
        
        generate_response = requests.post(
            generate_url,
            headers=generate_headers,
            json=generate_payload,
            timeout=120  # 2 minute timeout for API call
        )
        
        if generate_response.status_code != 200:
            logger.error(f"HeyGen v2 video generation failed: {generate_response.status_code} - {generate_response.text}")
            raise HTTPException(
                status_code=generate_response.status_code,
                detail=f"Failed to generate HeyGen video (v2): {generate_response.text}"
            )
        
        heygen_response = generate_response.json()
        logger.info(f"HeyGen v2 video generation initiated successfully: {heygen_response}")
        
        # Return HeyGen response with compatibility notes
        return JSONResponse(status_code=200, content={
            **heygen_response,
            "implementation": {
                "endpoint": "https://api.heygen.com/v2/video/generate",
                "avatar_iv_enabled": True,
                "talking_photo_id": talking_photo_id,
                "talking_style": "expressive",
                "super_resolution_enabled": request.super_resolution if request.super_resolution is not None else True,
                "ignored_legacy_fields_if_provided": [
                    "fit",
                    "custom_motion_prompt",
                    "enhance_custom_motion_prompt"
                ]
            }
        })
    
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error during HeyGen API call: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to connect to HeyGen API: {str(e)}"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during HeyGen Avatar IV generation: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate Avatar IV video: {str(e)}"
        )

@app.get("/telegram-file/bot{bot_token}/{file_path:path}")
async def telegram_file_to_s3(bot_token: str, file_path: str):
    """
    Download a Telegram file using bot token and file path, upload to S3 under audio/, and return the public URL.

    Telegram download URL format:
    https://api.telegram.org/file/bot<token>/<file_path>
    """
    logger.info(f"Telegram file request received - Path: {file_path}")

    try:
        # Validate required AWS environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )

        # Check if S3 client is available
        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )

        # Build Telegram file download URL
        telegram_url = f"https://api.telegram.org/file/bot{bot_token}/{file_path}"
        logger.info(f"Downloading Telegram file from: {telegram_url}")

        try:
            download_response = requests.get(
                telegram_url,
                timeout=300,
                stream=True
            )
            download_response.raise_for_status()
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to download Telegram file: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to download Telegram file: {str(e)}"
            )

        file_content = download_response.content
        if not file_content:
            logger.error("Downloaded Telegram file is empty")
            raise HTTPException(
                status_code=400,
                detail="Downloaded Telegram file is empty"
            )

        # Determine content type
        content_type = download_response.headers.get('Content-Type', 'application/octet-stream')
        if ';' in content_type:
            content_type = content_type.split(';')[0].strip()

        # Determine filename from file_path
        filename = os.path.basename(file_path)
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"telegram_audio_{timestamp}_{unique_id}"

        # Ensure audio/ prefix
        s3_key = filename
        if not s3_key.startswith("audio/"):
            s3_key = f"audio/{s3_key}"

        logger.info(f"Uploading Telegram file to S3 - Bucket: {S3_BUCKET_NAME}, Key: {s3_key}")

        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=s3_key,
            Body=file_content,
            ContentType=content_type,
            ACL='public-read'
        )

        logger.info(f"Telegram file uploaded to S3 successfully: {s3_key}")

        # Generate public URL
        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
        logger.info(f"Generated public URL: {public_url}")

        response_data = {
            "success": True,
            "message": "Telegram file downloaded and uploaded successfully",
            "public_url": public_url,
            "filename": filename,
            "s3_key": s3_key,
            "content_type": content_type,
            "uploaded_at": datetime.now().isoformat()
        }

        return JSONResponse(status_code=200, content=response_data)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during Telegram file upload: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process Telegram file: {str(e)}"
        )

def extract_google_drive_file_id(file_id_or_url: str) -> str:
    """
    Extract Google Drive file ID from various URL formats or return the ID if already provided.
    
    Supported formats:
    - File ID only: "1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H"
    - Share link: "https://drive.google.com/file/d/1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H/view"
    - Download link: "https://drive.google.com/uc?id=1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H&export=download"
    - Open link: "https://drive.google.com/open?id=1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H"
    
    Args:
        file_id_or_url: Either a file ID or a full Google Drive URL
        
    Returns:
        The extracted file ID
        
    Raises:
        ValueError: If the file ID cannot be extracted
    """
    
    # Strip whitespace
    file_id_or_url = file_id_or_url.strip()
    
    # Check if it's a URL
    if file_id_or_url.startswith('http://') or file_id_or_url.startswith('https://'):
        parsed_url = urlparse(file_id_or_url)
        
        # Check if it's a Google Drive URL
        if 'drive.google.com' not in parsed_url.netloc and 'docs.google.com' not in parsed_url.netloc:
            raise ValueError(f"URL is not a Google Drive link: {file_id_or_url}")
        
        # Try to extract from query parameter 'id'
        # Format: https://drive.google.com/uc?id=FILE_ID&export=download
        # Format: https://drive.google.com/open?id=FILE_ID
        query_params = parse_qs(parsed_url.query)
        if 'id' in query_params:
            return query_params['id'][0]
        
        # Try to extract from path
        # Format: https://drive.google.com/file/d/FILE_ID/view
        # Format: https://drive.google.com/file/d/FILE_ID/edit
        path_match = re.search(r'/file/d/([a-zA-Z0-9_-]+)', parsed_url.path)
        if path_match:
            return path_match.group(1)
        
        # Try to extract from folders path
        # Format: https://drive.google.com/drive/folders/FILE_ID
        folder_match = re.search(r'/folders/([a-zA-Z0-9_-]+)', parsed_url.path)
        if folder_match:
            return folder_match.group(1)
        
        # Try to extract ID from any part of the URL using regex
        # Google Drive IDs are typically 28-33 characters with letters, numbers, underscores, and hyphens
        id_match = re.search(r'([a-zA-Z0-9_-]{25,})', file_id_or_url)
        if id_match:
            return id_match.group(1)
        
        raise ValueError(f"Could not extract file ID from URL: {file_id_or_url}")
    
    # Check if it looks like a valid Google Drive file ID
    # Google Drive IDs are typically alphanumeric with underscores and hyphens, 25+ characters
    if re.match(r'^[a-zA-Z0-9_-]{10,}$', file_id_or_url):
        return file_id_or_url
    
    raise ValueError(f"Invalid Google Drive file ID or URL: {file_id_or_url}")

@app.post("/google-drive-to-s3")
async def google_drive_to_s3(request: GoogleDriveToS3Request):
    """
    Download a publicly accessible file from Google Drive and upload it to AWS S3.
    
    This endpoint:
    1. Takes a Google Drive file ID or URL (various formats supported)
    2. Downloads the file with retry logic to handle large files and confirmation pages
    3. Uploads to S3 at the specified folder path
    4. Returns the public S3 URL
    
    Supported URL formats:
    - File ID only: "1Rs1rijuqphRVdxT6WDcPG86x9IMKG81H"
    - Share link: "https://drive.google.com/file/d/FILE_ID/view"
    - Download link: "https://drive.google.com/uc?id=FILE_ID&export=download"
    - Open link: "https://drive.google.com/open?id=FILE_ID"
    
    Note: The Google Drive file must be shared as "Anyone with the link can view"
    
    Args:
        request: GoogleDriveToS3Request containing:
            - file_id: Google Drive file ID or URL (ID will be extracted automatically)
            - folder_path: S3 folder path (e.g., "videos/", "images/")
            - filename: Optional filename (auto-detected if not provided)
    
    Returns:
        JSON response with the public S3 URL of the uploaded file
    """
    logger.info(f"Google Drive to S3 request received - Input: {request.file_id}, Folder: {request.folder_path}")
    
    try:
        # Extract file ID from URL or use as-is if already an ID
        try:
            file_id = extract_google_drive_file_id(request.file_id)
            logger.info(f"Extracted Google Drive file ID: {file_id}")
        except ValueError as e:
            logger.error(f"Failed to extract file ID: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=str(e)
            )
        
        # Validate required AWS environment variables
        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]):
            logger.error("Missing required AWS environment variables")
            raise HTTPException(
                status_code=500,
                detail="AWS credentials or bucket name not configured"
            )
        
        # Check if S3 client is available
        if not s3_client:
            logger.error("S3 client not initialized")
            raise HTTPException(
                status_code=500,
                detail="S3 client not available - check AWS configuration"
            )
        
        # Normalize folder path (ensure it ends with / and doesn't start with /)
        folder_path = request.folder_path.strip('/')
        if folder_path:
            folder_path = folder_path + '/'
        
        # Google Drive download URL for publicly shared files
        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        
        # Configure retry settings
        max_retries = 3
        retry_delay = 2  # seconds
        
        file_content = None
        filename = request.filename
        content_type = "application/octet-stream"
        
        # Create a session to handle cookies (needed for large file confirmation)
        session = requests.Session()
        
        for attempt in range(max_retries):
            try:
                logger.info(f"Download attempt {attempt + 1}/{max_retries} for file ID: {file_id}")
                
                # Initial request
                response = session.get(
                    download_url,
                    stream=True,
                    timeout=30,
                    headers={
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                    }
                )
                response.raise_for_status()
                
                # Check if we got a confirmation page (for large files)
                # Google Drive returns HTML with a confirmation token for large files
                content_type_header = response.headers.get('Content-Type', '')
                
                if 'text/html' in content_type_header:
                    logger.info("Received HTML response, checking for download confirmation...")
                    
                    # Look for the confirmation token in the response
                    html_content = response.text

                    # Collect possible confirmation tokens from modern/legacy Google Drive pages.
                    confirm_candidates = []

                    # Pattern 1: URL parameter style (e.g. ...&confirm=xxxx...)
                    confirm_param_matches = re.findall(r'[?&]confirm=([0-9A-Za-z_-]+)', html_content)
                    confirm_candidates.extend(confirm_param_matches)

                    # Pattern 2: hidden input style (e.g. <input name="confirm" value="t">)
                    confirm_input_matches = re.findall(r'name=["\']confirm["\']\s+value=["\']([0-9A-Za-z_-]+)["\']', html_content)
                    confirm_candidates.extend(confirm_input_matches)

                    # Pattern 3: download warning cookie set by Drive
                    for cookie_name, cookie_value in session.cookies.items():
                        if cookie_name.startswith("download_warning") and cookie_value:
                            confirm_candidates.append(cookie_value)

                    # De-duplicate while preserving order
                    unique_confirm_candidates = list(dict.fromkeys(confirm_candidates))

                    # Build fallback URLs:
                    # - try explicit tokens first
                    # - then common fallback token 't'
                    # - then modern driveusercontent endpoint
                    confirmation_urls = [
                        f"https://drive.google.com/uc?export=download&id={file_id}&confirm={token}"
                        for token in unique_confirm_candidates
                    ]
                    confirmation_urls.append(f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t")
                    confirmation_urls.append(f"https://drive.usercontent.google.com/download?id={file_id}&export=download&confirm=t")

                    if unique_confirm_candidates:
                        logger.info(f"Found {len(unique_confirm_candidates)} confirmation token candidate(s)")
                    else:
                        logger.warning("No explicit confirmation token found; trying fallback confirmation URLs")

                    # Try each confirmation URL until we get non-HTML content
                    confirmed_response = None
                    for confirmed_url in confirmation_urls:
                        logger.info(f"Trying confirmation URL: {confirmed_url}")
                        candidate_response = session.get(
                            confirmed_url,
                            stream=True,
                            timeout=300,  # 5 minutes for large files
                            headers={
                                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                            }
                        )
                        candidate_response.raise_for_status()

                        candidate_content_type = candidate_response.headers.get('Content-Type', '')
                        if 'text/html' not in candidate_content_type:
                            confirmed_response = candidate_response
                            logger.info("Confirmation URL returned downloadable content")
                            break

                    if confirmed_response is not None:
                        response = confirmed_response
                    else:
                        logger.warning("All confirmation URL attempts returned HTML")
                
                # Check content type again after potential redirect
                content_type_header = response.headers.get('Content-Type', 'application/octet-stream')
                if ';' in content_type_header:
                    content_type = content_type_header.split(';')[0].strip()
                else:
                    content_type = content_type_header
                
                # If we still have HTML, the file might not be publicly accessible
                if 'text/html' in content_type:
                    logger.error("File appears to not be publicly accessible or doesn't exist")
                    raise HTTPException(
                        status_code=400,
                        detail="Unable to download file. Please ensure the Google Drive file is shared as 'Anyone with the link can view'"
                    )
                
                # Try to get filename from Content-Disposition header if not provided
                if not filename:
                    content_disposition = response.headers.get('Content-Disposition', '')
                    if 'filename=' in content_disposition:
                        # Parse filename from header
                        filename_match = re.search(r'filename\*?=["\']?(?:UTF-8\'\')?([^"\';\s]+)["\']?', content_disposition)
                        if filename_match:
                            filename = unquote(filename_match.group(1))
                            logger.info(f"Extracted filename from header: {filename}")
                
                # Generate filename if still not available
                if not filename:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    unique_id = str(uuid.uuid4())[:8]
                    
                    # Determine extension based on content type
                    extension_map = {
                        'image/jpeg': '.jpg',
                        'image/png': '.png',
                        'image/gif': '.gif',
                        'image/webp': '.webp',
                        'image/svg+xml': '.svg',
                        'video/mp4': '.mp4',
                        'video/webm': '.webm',
                        'video/quicktime': '.mov',
                        'audio/mpeg': '.mp3',
                        'audio/wav': '.wav',
                        'audio/ogg': '.ogg',
                        'application/pdf': '.pdf',
                        'text/plain': '.txt',
                        'application/zip': '.zip',
                    }
                    extension = extension_map.get(content_type, '')
                    filename = f"gdrive_file_{timestamp}_{unique_id}{extension}"
                    logger.info(f"Generated filename: {filename}")
                
                # Download the file content
                logger.info("Downloading file content...")
                file_content = response.content
                content_length = len(file_content)
                
                if content_length == 0:
                    raise Exception("Downloaded file is empty")
                
                logger.info(f"File downloaded successfully. Size: {content_length} bytes ({content_length / (1024*1024):.2f} MB)")

                # Google Drive may return a generic/incorrect Content-Type for SVG files.
                # Trust the filename extension for SVG and force the correct MIME type.
                if filename and filename.lower().endswith('.svg'):
                    content_type = 'image/svg+xml'
                    logger.info("Overriding Content-Type to image/svg+xml based on .svg filename")
                
                # Success - break out of retry loop
                break
                
            except requests.exceptions.RequestException as e:
                logger.warning(f"Download attempt {attempt + 1} failed: {str(e)}")
                if attempt < max_retries - 1:
                    logger.info(f"Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    logger.error(f"All {max_retries} download attempts failed")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to download file from Google Drive after {max_retries} attempts: {str(e)}"
                    )
        
        if file_content is None:
            raise HTTPException(
                status_code=500,
                detail="Failed to download file content"
            )
        
        # Build the full S3 key (folder path + filename)
        s3_key = f"{folder_path}{filename}"
        
        logger.info(f"Uploading to S3 - Bucket: {S3_BUCKET_NAME}, Key: {s3_key}")
        
        # Upload to S3
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=s3_key,
            Body=file_content,
            ContentType=content_type,
            ContentDisposition="inline",
            Tagging="temporary=true",
            ACL='public-read'
        )
        
        logger.info(f"File uploaded to S3 successfully: {s3_key}")
        
        # Generate public URL
        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
        logger.info(f"Generated public URL: {public_url}")
        
        # Return success response
        response_data = {
            "success": True,
            "message": "File downloaded from Google Drive and uploaded to S3 successfully",
            "public_url": public_url,
            "filename": filename,
            "folder_path": folder_path,
            "s3_key": s3_key,
            "file_id": file_id,
            "original_input": request.file_id,
            "file_size_bytes": content_length,
            "content_type": content_type,
            "uploaded_at": datetime.now().isoformat()
        }
        
        logger.info("Google Drive to S3 transfer completed successfully")
        return JSONResponse(status_code=200, content=response_data)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during Google Drive to S3 transfer: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to transfer file from Google Drive to S3: {str(e)}"
        )

@app.post("/transcribe-audio")
async def transcribe_audio(
    request: TranscribeAudioRequest,
    authorization: Optional[str] = Header(None)
):
    """
    Transcribe an audio file using OpenAI's transcription API.

    Requires:
    - Authorization: Bearer <OpenAI API key>
    - Request body with audio_url
    """
    logger.info(f"Transcription request received - Audio URL: {request.audio_url}")

    try:
        # Validate authorization header
        if not authorization or not authorization.lower().startswith("bearer "):
            logger.error("Missing or invalid authorization header")
            raise HTTPException(
                status_code=401,
                detail="Authorization header with Bearer token is required"
            )

        # Extract API key from authorization header
        api_key = authorization.split(" ", 1)[1].strip()

        # Download audio file
        logger.info(f"Downloading audio from URL: {request.audio_url}")
        try:
            audio_response = requests.get(
                request.audio_url,
                timeout=300,
                stream=True
            )
            audio_response.raise_for_status()
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to download audio file: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to download audio file: {str(e)}"
            )

        audio_content = audio_response.content
        if not audio_content:
            logger.error("Downloaded audio file is empty")
            raise HTTPException(
                status_code=400,
                detail="Downloaded audio file is empty"
            )

        # Determine content type and filename
        content_type = audio_response.headers.get('Content-Type', 'application/octet-stream')
        if ';' in content_type:
            content_type = content_type.split(';')[0].strip()

        filename = None
        content_disposition = audio_response.headers.get('Content-Disposition', '')
        if 'filename=' in content_disposition:
            match = re.search(r'filename[*]?=["\']?([^"\';\s]+)["\']?', content_disposition)
            if match:
                filename = match.group(1)

        if not filename:
            parsed_url = urlparse(request.audio_url)
            path_filename = parsed_url.path.split('/')[-1]
            if path_filename:
                filename = unquote(path_filename)
            else:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                unique_id = str(uuid.uuid4())[:8]
                filename = f"audio_{timestamp}_{unique_id}"

        # Call OpenAI transcription API
        openai_url = "https://api.openai.com/v1/audio/transcriptions"
        openai_headers = {
            "Authorization": f"Bearer {api_key}"
        }

        files = {
            "file": (filename, audio_content, content_type)
        }
        data = {
            "model": "whisper-1"
        }

        logger.info("Calling OpenAI transcription API")
        openai_response = requests.post(
            openai_url,
            headers=openai_headers,
            files=files,
            data=data,
            timeout=300
        )

        if openai_response.status_code != 200:
            logger.error(f"OpenAI API error: {openai_response.status_code} - {openai_response.text}")
            raise HTTPException(
                status_code=openai_response.status_code,
                detail=f"OpenAI API error: {openai_response.text}"
            )

        response_json = openai_response.json()
        transcript_text = response_json.get("text")

        if not transcript_text:
            logger.error(f"No transcript returned from OpenAI: {response_json}")
            raise HTTPException(
                status_code=500,
                detail="No transcript returned from OpenAI"
            )

        response_data = {
            "success": True,
            "message": "Transcription completed successfully",
            "transcript": transcript_text,
            "filename": filename,
            "content_type": content_type,
            "transcribed_at": datetime.now().isoformat()
        }

        return JSONResponse(status_code=200, content=response_data)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during transcription: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to transcribe audio: {str(e)}"
        )

@app.post("/analyze-images")
async def analyze_images(
    request: AnalyzeImagesRequest,
    authorization: Optional[str] = Header(None)
):
    """
    Analyze multiple images using a single OpenAI request.

    Request body:
    - prompt: Analysis prompt
    - image_urls: Comma-separated image URLs
    - model: OpenAI model to use
    - reasoning_effort: Optional reasoning effort (empty string becomes "none")
    """
    logger.info("Analyze images request received")

    try:
        request_started_at = time.perf_counter()
        # Validate authorization header
        if not authorization or not authorization.lower().startswith("bearer "):
            logger.error("Missing or invalid authorization header")
            raise HTTPException(
                status_code=401,
                detail="Authorization header with Bearer token is required"
            )

        # Extract API key from authorization header
        api_key = authorization.split(" ", 1)[1].strip()

        # Parse and clean URL list
        urls = [
            u.strip().replace(" ", "%20")
            for u in request.image_urls.split(",")
            if u.strip()
        ]

        if len(urls) == 0:
            logger.error("No valid image URLs found in image_urls")
            raise HTTPException(
                status_code=400,
                detail="image_urls must contain at least one valid URL"
            )

        content_array = [
            {"type": "text", "text": request.prompt}
        ]

        total_original_bytes = 0
        total_optimized_bytes = 0

        def download_single_image(index: int, url: str):
            logger.info(f"Downloading image {index + 1}/{len(urls)}")
            try:
                image_response = requests.get(
                    url,
                    timeout=180,
                    stream=True
                )
                image_response.raise_for_status()
            except requests.exceptions.RequestException as e:
                logger.error(f"Failed to download image at index {index}: {str(e)}")
                raise HTTPException(
                    status_code=400,
                    detail=f"Failed to download image at index {index}: {str(e)}"
                )

            image_bytes = image_response.content
            if not image_bytes:
                logger.error(f"Downloaded image at index {index} is empty")
                raise HTTPException(
                    status_code=400,
                    detail=f"Downloaded image at index {index} is empty"
                )

            content_type = image_response.headers.get("Content-Type", "").split(";")[0].strip()
            if not content_type:
                parsed_url = urlparse(url)
                guessed_content_type, _ = mimetypes.guess_type(parsed_url.path)
                content_type = guessed_content_type or "image/jpeg"

            return {
                "index": index,
                "url": url,
                "image_bytes": image_bytes,
                "content_type": content_type
            }

        # Download images in parallel to reduce total wall-clock time.
        download_started_at = time.perf_counter()
        max_workers = min(8, len(urls))
        logger.info(f"Starting parallel image downloads with {max_workers} workers")
        downloaded_images = {}
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {
                executor.submit(download_single_image, index, url): index
                for index, url in enumerate(urls)
            }
            for future in as_completed(future_map):
                result = future.result()
                downloaded_images[result["index"]] = result
        download_elapsed_seconds = time.perf_counter() - download_started_at

        def process_single_image(image_item: dict):
            index = image_item["index"]
            image_bytes = image_item["image_bytes"]
            content_type = image_item["content_type"]

            original_size = len(image_bytes)

            # Optimize image bytes to reduce payload size/timeouts.
            optimized_bytes = image_bytes
            optimized_content_type = content_type
            try:
                with Image.open(BytesIO(image_bytes)) as img:
                    # Honor EXIF orientation before resizing.
                    img = ImageOps.exif_transpose(img)

                    max_dimension = 1568
                    if max(img.size) > max_dimension:
                        img.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)

                    output = BytesIO()
                    has_alpha = img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info)

                    if has_alpha:
                        img.save(output, format="PNG", optimize=True, compress_level=9)
                        optimized_content_type = "image/png"
                    else:
                        if img.mode not in ("RGB", "L"):
                            img = img.convert("RGB")
                        img.save(output, format="JPEG", quality=82, optimize=True, progressive=True)
                        optimized_content_type = "image/jpeg"

                    candidate_bytes = output.getvalue()
                    if candidate_bytes and len(candidate_bytes) < len(image_bytes):
                        optimized_bytes = candidate_bytes
            except Exception as e:
                logger.warning(f"Image optimization failed at index {index}, using original bytes: {str(e)}")

            optimized_size = len(optimized_bytes)
            logger.info(
                f"Image {index} size bytes - original: {original_size}, optimized: {optimized_size}, content-type: {optimized_content_type}"
            )

            encoded_image = base64.b64encode(optimized_bytes).decode("utf-8")
            data_url = f"data:{optimized_content_type};base64,{encoded_image}"

            return {
                "index": index,
                "original_size": original_size,
                "optimized_size": optimized_size,
                "label": {
                    "type": "text",
                    "text": f"Image Index {index}:"
                },
                "image": {
                    "type": "image_url",
                    "image_url": {
                        "url": data_url
                    }
                }
            }

        # Process images in parallel, then build content array in original order.
        processing_started_at = time.perf_counter()
        process_workers = min(6, len(urls))
        logger.info(f"Starting parallel image processing with {process_workers} workers")
        processed_images = {}
        with ThreadPoolExecutor(max_workers=process_workers) as executor:
            process_future_map = {
                executor.submit(process_single_image, image_item): image_item["index"]
                for image_item in downloaded_images.values()
            }
            for future in as_completed(process_future_map):
                result = future.result()
                processed_images[result["index"]] = result
        processing_elapsed_seconds = time.perf_counter() - processing_started_at

        for index in range(len(urls)):
            processed = processed_images[index]
            total_original_bytes += processed["original_size"]
            total_optimized_bytes += processed["optimized_size"]
            content_array.append(processed["label"])
            content_array.append(processed["image"])

        reasoning_effort = request.reasoning_effort if request.reasoning_effort != "" else "none"

        openai_payload = {
            "model": request.model,
            "messages": [
                {
                    "role": "user",
                    "content": content_array
                }
            ],
            "reasoning_effort": reasoning_effort
        }

        openai_headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        logger.info(f"Sending image analysis request to OpenAI for {len(urls)} images")
        openai_started_at = time.perf_counter()
        openai_response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=openai_headers,
            json=openai_payload,
            timeout=600
        )
        openai_elapsed_seconds = time.perf_counter() - openai_started_at

        if openai_response.status_code != 200:
            logger.error(f"OpenAI API error: {openai_response.status_code} - {openai_response.text}")
            raise HTTPException(
                status_code=openai_response.status_code,
                detail=f"OpenAI API error: {openai_response.text}"
            )

        response_json = openai_response.json()
        analysis_text = ""
        if response_json.get("choices") and len(response_json["choices"]) > 0:
            analysis_text = response_json["choices"][0].get("message", {}).get("content", "")

        total_elapsed_seconds = time.perf_counter() - request_started_at
        logger.info(
            f"Analyze images timing - total: {total_elapsed_seconds:.2f}s, "
            f"download: {download_elapsed_seconds:.2f}s, "
            f"processing: {processing_elapsed_seconds:.2f}s, "
            f"openai: {openai_elapsed_seconds:.2f}s, "
            f"image_count: {len(urls)}"
        )

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Image analysis completed successfully",
                "image_count": len(urls),
                "model": request.model,
                "reasoning_effort": reasoning_effort,
                "analysis": analysis_text,
                "raw_response": response_json,
                "payload_bytes": {
                    "original_total": total_original_bytes,
                    "optimized_total": total_optimized_bytes,
                    "saved_bytes": max(total_original_bytes - total_optimized_bytes, 0)
                },
                "timings_seconds": {
                    "download": round(download_elapsed_seconds, 3),
                    "processing": round(processing_elapsed_seconds, 3),
                    "openai": round(openai_elapsed_seconds, 3),
                    "total": round(total_elapsed_seconds, 3)
                },
                "analyzed_at": datetime.now().isoformat()
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during image analysis: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to analyze images: {str(e)}"
        )

@app.post("/markdown-to-docx")
async def markdown_to_docx(request: MarkdownToDocxRequest):
    """
    Convert markdown text to a .docx Word document and upload to S3 under documents/.

    Supports headers (h1-h6), paragraphs, bold/italic/inline-code, links,
    ordered/unordered lists, tables, blockquotes, code blocks, and horizontal rules.
    """
    logger.info(f"Markdown-to-docx request - filename: {request.filename}, length: {len(request.text)}")

    try:
        import markdown as md_lib
        from bs4 import BeautifulSoup, NavigableString
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not all([AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME]) or not s3_client:
            raise HTTPException(status_code=500, detail="AWS/S3 not configured")

        # 1. Markdown -> HTML
        html = md_lib.markdown(
            request.text,
            extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
        )
        soup = BeautifulSoup(f"<div>{html}</div>", "html.parser")

        # 2. Build the docx
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        def add_inline(paragraph, node, bold=False, italic=False, code=False, link=None):
            """Recursively add inline runs to a paragraph."""
            if isinstance(node, NavigableString):
                text = str(node)
                if not text:
                    return
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Consolas"
                if link:
                    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
                    run.underline = True
                return

            name = node.name
            new_bold = bold or name in ("strong", "b")
            new_italic = italic or name in ("em", "i")
            new_code = code or name == "code"
            new_link = link or (node.get("href") if name == "a" else None)

            if name == "br":
                paragraph.add_run().add_break()
                return

            for child in node.children:
                add_inline(paragraph, child, new_bold, new_italic, new_code, new_link)

        def add_list(list_node, ordered, level=0):
            style_name = "List Number" if ordered else "List Bullet"
            for li in list_node.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style_name)
                p.paragraph_format.left_indent = Pt(18 * (level + 1))
                # Inline content of the <li>, excluding nested lists
                for child in li.children:
                    if getattr(child, "name", None) in ("ul", "ol"):
                        continue
                    add_inline(p, child)
                # Recurse into nested lists
                for nested in li.find_all(["ul", "ol"], recursive=False):
                    add_list(nested, nested.name == "ol", level + 1)

        def add_table(table_node):
            rows = table_node.find_all("tr")
            if not rows:
                return
            cols = max(len(r.find_all(["td", "th"])) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Light Grid Accent 1"
            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["td", "th"])
                for c_idx, cell in enumerate(cells):
                    docx_cell = tbl.rows[r_idx].cells[c_idx]
                    docx_cell.text = ""
                    p = docx_cell.paragraphs[0]
                    is_header = cell.name == "th"
                    for child in cell.children:
                        add_inline(p, child, bold=is_header)

        root = soup.find("div")
        for el in root.children:
            if isinstance(el, NavigableString):
                txt = str(el).strip()
                if txt:
                    doc.add_paragraph(txt)
                continue

            tag = el.name
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(tag[1])
                p = doc.add_heading(level=level if level <= 9 else 9)
                for child in el.children:
                    add_inline(p, child)
            elif tag == "p":
                p = doc.add_paragraph()
                for child in el.children:
                    add_inline(p, child)
            elif tag == "ul":
                add_list(el, ordered=False)
            elif tag == "ol":
                add_list(el, ordered=True)
            elif tag == "table":
                add_table(el)
            elif tag == "blockquote":
                for child_p in el.find_all("p"):
                    p = doc.add_paragraph(style="Intense Quote")
                    for child in child_p.children:
                        add_inline(p, child)
            elif tag == "pre":
                code_text = el.get_text()
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif tag == "hr":
                doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                for child in el.children:
                    add_inline(p, child)

        # 3. Serialize to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()

        # 4. Build S3 key under documents/
        if request.filename:
            base = request.filename
            if base.endswith(".docx"):
                base = base[:-5]
            base = base.lstrip("/")
            if base.startswith("documents/"):
                base = base[len("documents/"):]
            key = f"documents/{base}.docx"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            key = f"documents/markdown_{timestamp}_{unique_id}.docx"

        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=key,
            Body=docx_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ACL="public-read",
        )

        public_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{key}"
        logger.info(f"Uploaded docx to S3: {key}")

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Markdown converted and uploaded successfully",
                "filename": key,
                "public_url": public_url,
                "bucket": S3_BUCKET_NAME,
                "region": AWS_REGION,
                "key": key,
                "size_bytes": len(docx_bytes),
                "uploaded_at": datetime.now().isoformat(),
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error converting markdown to docx: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to convert markdown: {str(e)}")

@app.post("/list-sharepoint-subfolders")
async def list_sharepoint_subfolders(request: ListSharepointSubfoldersRequest):
    """
    List subfolders (and optionally files) inside a SharePoint folder via
    Microsoft Graph using app-only auth.

    Body:
        sharepoint_url: SharePoint URL of the parent folder
        recursive: If true, walk the folder tree recursively
        include_files: If true, include files in the response alongside folders
    """
    logger.info(
        f"List-sharepoint-subfolders request - parent: {request.sharepoint_url}, "
        f"recursive: {request.recursive}, include_files: {request.include_files}"
    )

    GRAPH = "https://graph.microsoft.com/v1.0"
    drive_id, folder_id = _resolve_sharepoint_folder(request.sharepoint_url)
    auth_headers = {"Authorization": f"Bearer {_get_graph_app_token()}"}

    def _fetch_children(parent_id):
        out = []
        url = (
            f"{GRAPH}/drives/{drive_id}/items/{parent_id}/children"
            f"?$top=200&$select=id,name,webUrl,size,folder,file,parentReference"
        )
        while url:
            r = requests.get(url, headers=auth_headers, timeout=30)
            if r.status_code != 200:
                raise HTTPException(
                    status_code=502,
                    detail=f"Failed to list folder contents: {r.text}",
                )
            j = r.json()
            out.extend(j.get("value", []))
            url = j.get("@odata.nextLink")
        return out

    def _shape(item, path):
        is_folder = "folder" in item
        return {
            "id": item.get("id"),
            "name": item.get("name"),
            "is_folder": is_folder,
            "path": path,
            "web_url": item.get("webUrl"),
            "size": item.get("size"),
            "child_count": (item.get("folder") or {}).get("childCount") if is_folder else None,
        }

    results = []
    if not request.recursive:
        for item in _fetch_children(folder_id):
            if "folder" in item or request.include_files:
                results.append(_shape(item, item.get("name", "")))
    else:
        # BFS walk; track relative paths from the requested folder.
        queue = [(folder_id, "")]
        while queue:
            current_id, prefix = queue.pop(0)
            for item in _fetch_children(current_id):
                rel_path = (
                    f"{prefix}/{item.get('name', '')}".lstrip("/")
                    if prefix
                    else item.get("name", "")
                )
                if "folder" in item:
                    results.append(_shape(item, rel_path))
                    queue.append((item["id"], rel_path))
                elif request.include_files:
                    results.append(_shape(item, rel_path))

    return JSONResponse(
        status_code=200,
        content={
            "success": True,
            "drive_id": drive_id,
            "parent_folder_id": folder_id,
            "count": len(results),
            "items": results,
            "listed_at": datetime.now().isoformat(),
        },
    )


@app.post("/create-sharepoint-folder")
async def create_sharepoint_folder(request: CreateSharepointFolderRequest):
    """
    Create a folder inside a SharePoint folder via Microsoft Graph using
    app-only auth.

    Body:
        sharepoint_url: SharePoint URL of the parent folder
        folder_name: Name of the new folder to create
        conflict_behavior: "rename" (default), "replace", or "fail"
    """
    logger.info(
        f"Create-sharepoint-folder request - parent: {request.sharepoint_url}, "
        f"name: {request.folder_name}"
    )

    folder_name = (request.folder_name or "").strip()
    if not folder_name:
        raise HTTPException(status_code=400, detail="folder_name is required")
    if any(c in folder_name for c in '\\/:*?"<>|'):
        raise HTTPException(
            status_code=400,
            detail='folder_name contains invalid characters: \\/:*?"<>|',
        )

    conflict = request.conflict_behavior or "rename"
    if conflict not in ("rename", "replace", "fail"):
        raise HTTPException(
            status_code=400,
            detail="conflict_behavior must be 'rename', 'replace', or 'fail'",
        )

    GRAPH = "https://graph.microsoft.com/v1.0"
    drive_id, parent_id = _resolve_sharepoint_folder(request.sharepoint_url)
    auth_headers = {"Authorization": f"Bearer {_get_graph_app_token()}"}

    create_resp = requests.post(
        f"{GRAPH}/drives/{drive_id}/items/{parent_id}/children",
        headers={**auth_headers, "Content-Type": "application/json"},
        json={
            "name": folder_name,
            "folder": {},
            "@microsoft.graph.conflictBehavior": conflict,
        },
        timeout=30,
    )
    if create_resp.status_code not in (200, 201):
        logger.error(
            f"Folder creation failed: {create_resp.status_code} {create_resp.text}"
        )
        raise HTTPException(
            status_code=502,
            detail=f"Failed to create SharePoint folder: {create_resp.text}",
        )

    item = create_resp.json()
    logger.info(
        f"Created SharePoint folder id={item.get('id')} name={item.get('name')}"
    )

    return JSONResponse(
        status_code=200,
        content={
            "success": True,
            "message": "Folder created successfully",
            "item_id": item.get("id"),
            "name": item.get("name"),
            "web_url": item.get("webUrl"),
            "drive_id": drive_id,
            "parent_folder_id": parent_id,
            "created_at": datetime.now().isoformat(),
        },
    )


@app.post("/upload-to-sharepoint")
async def upload_to_sharepoint(request: UploadToSharepointRequest):
    """
    Download a file from a public URL and upload it to a SharePoint folder
    via Microsoft Graph using app-only (client-credentials) authentication.

    Credentials come from the env: SHAREPOINT_TENANT_ID, SHAREPOINT_CLIENT_ID,
    SHAREPOINT_CLIENT_SECRET. The app registration must have the application
    permission `Sites.ReadWrite.All` granted with admin consent.

    Body:
        file_url: Publicly accessible URL of the file to upload
        sharepoint_url: SharePoint folder URL (browser-copied URL works)
        filename: Optional override for the uploaded filename
    """
    logger.info(
        f"Upload-to-sharepoint request - file_url: {request.file_url}, "
        f"sharepoint_url: {request.sharepoint_url}, filename: {request.filename}"
    )

    GRAPH = "https://graph.microsoft.com/v1.0"
    access_token = _get_graph_app_token()
    auth_headers = {"Authorization": f"Bearer {access_token}"}

    try:
        # 1. Parse the SharePoint URL into hostname + site path + folder path.
        parsed = urlparse(request.sharepoint_url)
        hostname = parsed.netloc
        path = unquote(parsed.path).lstrip("/")
        path = re.sub(r"^:[a-z]:/[a-z]/", "", path)
        parts = [p for p in path.split("/") if p]
        if len(parts) < 2 or parts[0] not in ("sites", "teams"):
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Could not parse SharePoint URL. Expected a path containing "
                    f"'sites/<name>' or 'teams/<name>'. Got: {request.sharepoint_url}"
                ),
            )
        site_segment = f"{parts[0]}/{parts[1]}"
        remaining = parts[2:]  # e.g. ["Shared Documents", "Folder", "Subfolder"]

        # 2. Resolve site -> siteId.
        site_resp = requests.get(
            f"{GRAPH}/sites/{hostname}:/{site_segment}",
            headers=auth_headers,
            timeout=30,
        )
        if site_resp.status_code != 200:
            raise HTTPException(
                status_code=502,
                detail=f"Failed to resolve SharePoint site: {site_resp.text}",
            )
        site_id = site_resp.json().get("id")

        # 3. Pick the drive whose name/webUrl matches the library segment;
        #    fall back to the site's default drive.
        library_name = remaining[0] if remaining else None
        folder_path_parts = remaining[1:] if remaining else []

        drives_resp = requests.get(
            f"{GRAPH}/sites/{site_id}/drives", headers=auth_headers, timeout=30
        )
        candidate_drives = []
        seen = set()
        if drives_resp.status_code == 200:
            for d in drives_resp.json().get("value", []):
                if d.get("id") and d["id"] not in seen:
                    candidate_drives.append(d)
                    seen.add(d["id"])
        default_resp = requests.get(
            f"{GRAPH}/sites/{site_id}/drive", headers=auth_headers, timeout=30
        )
        if default_resp.status_code == 200:
            d = default_resp.json()
            if d.get("id") and d["id"] not in seen:
                candidate_drives.append(d)
                seen.add(d["id"])
        if not candidate_drives:
            raise HTTPException(
                status_code=502, detail="No SharePoint drives accessible for this site"
            )
        if library_name:
            def _score(d):
                name = d.get("name", "")
                web = d.get("webUrl", "").rstrip("/")
                if name == library_name:
                    return 0
                if web.endswith("/" + library_name.replace(" ", "%20")):
                    return 1
                return 2
            candidate_drives.sort(key=_score)

        def _walk(drive, segments):
            drv_id = drive.get("id")
            root_resp_local = requests.get(
                f"{GRAPH}/drives/{drv_id}/root", headers=auth_headers, timeout=30
            )
            if root_resp_local.status_code != 200:
                return None, f"root fetch failed: {root_resp_local.text}"
            cursor = root_resp_local.json().get("id")
            for segment in segments:
                children_url = (
                    f"{GRAPH}/drives/{drv_id}/items/{cursor}/children"
                    f"?$top=200&$select=id,name,folder"
                )
                found_child = None
                seen_names = []
                while children_url:
                    ch_resp = requests.get(children_url, headers=auth_headers, timeout=30)
                    if ch_resp.status_code != 200:
                        return None, f"children fetch failed: {ch_resp.text}"
                    ch_json = ch_resp.json()
                    for child in ch_json.get("value", []):
                        seen_names.append(child.get("name", ""))
                        if child.get("name", "").lower() == segment.lower():
                            found_child = child
                            break
                    if found_child:
                        break
                    children_url = ch_json.get("@odata.nextLink")
                if not found_child:
                    return None, (
                        f"segment '{segment}' not found; available: {seen_names[:20]}"
                    )
                cursor = found_child["id"]
            return cursor, None

        attempt_segments = [folder_path_parts]
        if folder_path_parts != remaining:
            attempt_segments.append(remaining)

        walked_drive = None
        folder_id = None
        diagnostics = []
        for d in candidate_drives:
            for segs in attempt_segments:
                fid, err = _walk(d, segs)
                if fid:
                    walked_drive = d
                    folder_id = fid
                    break
                diagnostics.append(
                    f"drive '{d.get('name')}' ({d.get('id')}) segs={segs}: {err}"
                )
            if folder_id:
                break

        if not folder_id:
            raise HTTPException(
                status_code=502,
                detail=(
                    f"Failed to resolve folder path '{'/'.join(remaining)}'. "
                    f"Diagnostics: {diagnostics}"
                ),
            )

        drive_id = walked_drive.get("id")

        # 3. Download the source file.
        file_resp = requests.get(request.file_url, stream=True, timeout=120)
        if file_resp.status_code != 200:
            raise HTTPException(
                status_code=502,
                detail=f"Failed to download source file: HTTP {file_resp.status_code}",
            )
        file_bytes = file_resp.content
        file_size = len(file_bytes)

        # Determine filename
        if request.filename:
            upload_name = request.filename
        else:
            parsed = urlparse(request.file_url)
            upload_name = unquote(os.path.basename(parsed.path)) or f"upload_{uuid.uuid4().hex[:8]}"
        # Make safe-ish for SharePoint
        upload_name = upload_name.replace("/", "_").replace("\\", "_").strip()
        if not upload_name:
            upload_name = f"upload_{uuid.uuid4().hex[:8]}"

        logger.info(
            f"Resolved folder driveId={drive_id} itemId={folder_id}; uploading {upload_name} ({file_size} bytes)"
        )

        # 4. Create an upload session and PUT the file in chunks.
        session_url = (
            f"{GRAPH}/drives/{drive_id}/items/{folder_id}:/"
            f"{requests.utils.quote(upload_name)}:/createUploadSession"
        )
        session_body = {
            "item": {
                "@microsoft.graph.conflictBehavior": "rename",
                "name": upload_name,
            }
        }
        session_resp = requests.post(
            session_url,
            headers={**auth_headers, "Content-Type": "application/json"},
            json=session_body,
            timeout=30,
        )
        if session_resp.status_code not in (200, 201):
            logger.error(
                f"Failed to create upload session: {session_resp.status_code} {session_resp.text}"
            )
            raise HTTPException(
                status_code=502,
                detail=f"Failed to create SharePoint upload session: {session_resp.text}",
            )
        upload_url = session_resp.json().get("uploadUrl")
        if not upload_url:
            raise HTTPException(status_code=502, detail="Upload session missing uploadUrl")

        # Chunked PUT. Graph requires multiples of 320 KiB for non-final chunks.
        CHUNK = 5 * 320 * 1024  # 1.6 MiB
        final_resp = None
        if file_size == 0:
            # Graph still expects a PUT for empty files
            final_resp = requests.put(
                upload_url,
                headers={"Content-Range": "bytes 0-0/0", "Content-Length": "0"},
                data=b"",
                timeout=60,
            )
        else:
            for start in range(0, file_size, CHUNK):
                end = min(start + CHUNK, file_size) - 1
                chunk = file_bytes[start : end + 1]
                headers = {
                    "Content-Length": str(len(chunk)),
                    "Content-Range": f"bytes {start}-{end}/{file_size}",
                }
                final_resp = requests.put(upload_url, headers=headers, data=chunk, timeout=300)
                if final_resp.status_code not in (200, 201, 202):
                    logger.error(
                        f"Chunk upload failed at {start}-{end}: "
                        f"{final_resp.status_code} {final_resp.text}"
                    )
                    # Best-effort cancel of the session
                    try:
                        requests.delete(upload_url, timeout=15)
                    except Exception:
                        pass
                    raise HTTPException(
                        status_code=502,
                        detail=f"SharePoint chunk upload failed: {final_resp.text}",
                    )

        if final_resp is None or final_resp.status_code not in (200, 201):
            raise HTTPException(
                status_code=502,
                detail=f"Upload did not complete successfully (status {getattr(final_resp, 'status_code', 'n/a')})",
            )

        uploaded_item = final_resp.json()
        logger.info(
            f"Successfully uploaded to SharePoint: id={uploaded_item.get('id')} "
            f"name={uploaded_item.get('name')}"
        )

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "File uploaded to SharePoint successfully",
                "item_id": uploaded_item.get("id"),
                "name": uploaded_item.get("name"),
                "web_url": uploaded_item.get("webUrl"),
                "size": uploaded_item.get("size"),
                "drive_id": drive_id,
                "parent_folder_id": folder_id,
                "uploaded_at": datetime.now().isoformat(),
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading to SharePoint: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to upload to SharePoint: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    logger.info("Health check endpoint accessed")
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.on_event("startup")
async def startup_event():
    logger.info("FastAPI application starting up...")
    logger.info(f"Environment check - AWS_REGION: {AWS_REGION}, S3_BUCKET: {S3_BUCKET_NAME}")
    logger.info(f"Initializing FastAPI app with AWS region: {AWS_REGION}")
    logger.info(f"S3 Bucket: {S3_BUCKET_NAME}")
    logger.info(f"AWS Access Key ID configured: {'Yes' if AWS_ACCESS_KEY_ID else 'No'}")
    logger.info(f"AWS Secret Access Key configured: {'Yes' if AWS_SECRET_ACCESS_KEY else 'No'}")
    logger.info(f"Assembly AI API Key configured: {'Yes' if ASSEMBLY_AI_API_KEY else 'No'}")
    logger.info(f"Environment variables loaded from .env file")

@app.on_event("shutdown")
async def shutdown_event():
    logger.info("FastAPI application shutting down...")

if __name__ == "__main__":
    import uvicorn
    logger.info("Starting uvicorn server...")
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
